"""
Report Generator Service (FR-12).
Generates PDF, HTML, and DOCX reports with embedded Matplotlib charts.
Uses WeasyPrint (PDF), Jinja2 (HTML), python-docx (DOCX).

Report types:
  R-01: Traffic Flow Report
  R-02: Resource Usage Report
  R-03: Active VPN Users Report
  R-04: SD-WAN SLA Report
  R-05: Traffic Inbound Report
  R-06: Traffic Internal Report
  R-07: Executive Summary Report
  R-08: All-in-One Report (R-01 through R-07 combined)
"""
from __future__ import annotations

import base64
import logging
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Optional

from weasyprint import HTML

from app.core.config import get_settings
from app.db.models import ReportJob

settings = get_settings()
logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path("reports/templates")

# Sites used for per-site traffic in R-01 and for SD-WAN SLA queries
DEFAULT_SITES = ["Site_FGT-DC", "Site_FGT-DRC", "Site_FGT_Office"]


async def generate_report(job: ReportJob) -> str:
    """
    Generate a report file for the given job.
    Returns the output file path.
    Raises Exception on failure.

    Report types:
      R-01: Traffic Flow Report
      R-02: Resource Usage Report
      R-03: Active VPN Users Report
      R-04: SD-WAN SLA Report
      R-05: Traffic Inbound Report
      R-06: Traffic Internal Report
      R-07: Executive Summary Report
      R-08: All-in-One Report (R-01 through R-07)
    """
    output_dir = Path("reports/output")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{job.id}.{job.output_format}"

    # Fetch data from OpenSearch
    gte_ms = int(job.time_range_start.timestamp() * 1000)
    lte_ms = int(job.time_range_end.timestamp() * 1000)

    context = await _build_report_context(job.report_type, gte_ms, lte_ms)
    context["report_title"] = _report_title(job.report_type)
    context["generated_at"] = datetime.now(timezone.utc).isoformat()
    context["generated_by"] = str(job.created_by)
    context["job_id"] = str(job.id)
    context["time_range"] = f"{job.time_range_start.isoformat()} — {job.time_range_end.isoformat()}"

    match job.output_format:
        case "html":
            _generate_html(context, output_path)
        case "pdf":
            _generate_pdf(context, output_path)
        case "docx":
            _generate_docx(context, output_path)
        case _:
            raise ValueError(f"Unsupported output format: {job.output_format}")

    logger.info(f"Report generated: {output_path}")
    return str(output_path)


def _report_title(report_type: str) -> str:
    mapping = {
        "R-01": "Traffic Flow Report",
        "R-02": "Resource Usage Report",
        "R-03": "Active VPN Users Report",
        "R-04": "SD-WAN SLA Report",
        "R-05": "Traffic Inbound Report",
        "R-06": "Traffic Internal Report",
        "R-07": "Executive Summary Report",
        "R-08": "All-in-One Network Observability Report",
    }
    return mapping.get(report_type, "NOD Report")


async def _build_report_context(
    report_type: str, gte_ms: int, lte_ms: int
) -> dict:
    """
    Fetch data from OpenSearch and build the report template context.
    Includes base64-encoded chart PNGs for embedding.
    """
    context: dict = {"charts": {}, "per_site_traffic": {}}
    from app.services.chart_renderer import render_bar_chart, render_timeseries_chart, render_vpn_bar_chart

    # ── R-01: Traffic Flow (enhanced) ────────────────────────────
    if report_type in ("R-01", "R-08"):
        try:
            from app.opensearch import appid as appid_qb

            # Top Applications bar chart
            top_apps = await appid_qb.top_applications(gte_ms=gte_ms, lte_ms=lte_ms, size=10)
            if top_apps:
                labels = [a["application"] for a in top_apps[:10]]
                values = [a["total_bytes"] for a in top_apps[:10]]
                png_bytes = render_bar_chart(
                    labels, values,
                    title="Top 10 Applications by Traffic Volume",
                    xlabel="Bytes",
                )
                context["charts"]["top_applications"] = base64.b64encode(png_bytes).decode()

            # Throughput timeline
            tp_data = await appid_qb.throughput_timeline(
                gte_ms=gte_ms, lte_ms=lte_ms, interval="15m"
            )
            if tp_data:
                png_bytes = render_timeseries_chart(
                    tp_data,
                    title="Throughput Over Time",
                    ylabel="Bytes",
                )
                context["charts"]["throughput"] = base64.b64encode(png_bytes).decode()

            # Total throughput
            total_bytes = await appid_qb.total_throughput(gte_ms=gte_ms, lte_ms=lte_ms)
            context["total_throughput_bytes"] = total_bytes
            context["top_applications"] = top_apps
        except Exception as exc:
            logger.error("R-01 base traffic data fetch failed: %s", exc, exc_info=True)

        # R-01 Enhanced: Top AS Orgs
        try:
            from app.opensearch import appid as appid_qb
            top_as_orgs = await appid_qb.top_dst_as_orgs(gte_ms=gte_ms, lte_ms=lte_ms, size=10)
            context["top_as_orgs"] = top_as_orgs
            if top_as_orgs:
                labels = [o["org_name"] for o in top_as_orgs[:10]]
                values = [o["total_bytes"] for o in top_as_orgs[:10]]
                png_bytes = render_bar_chart(
                    labels, values,
                    title="Top 10 Destination AS Organizations",
                    xlabel="Bytes",
                )
                context["charts"]["top_as_orgs"] = base64.b64encode(png_bytes).decode()
        except Exception as exc:
            logger.error("R-01 top_as_orgs fetch failed: %s", exc, exc_info=True)

        # R-01 Enhanced: Top Countries
        try:
            from app.opensearch import appid as appid_qb
            top_countries = await appid_qb.top_dst_as_countries(gte_ms=gte_ms, lte_ms=lte_ms, size=10)
            context["top_countries"] = top_countries
            if top_countries:
                labels = [c["country"] for c in top_countries[:10]]
                values = [c["total_bytes"] for c in top_countries[:10]]
                png_bytes = render_bar_chart(
                    labels, values,
                    title="Top 10 Destination Countries",
                    xlabel="Bytes",
                )
                context["charts"]["top_countries"] = base64.b64encode(png_bytes).decode()
        except Exception as exc:
            logger.error("R-01 top_countries fetch failed: %s", exc, exc_info=True)

        # R-01 Enhanced: Protocol Distribution
        try:
            from app.opensearch import appid as appid_qb
            protocol_dist = await appid_qb.protocol_distribution(gte_ms=gte_ms, lte_ms=lte_ms)
            context["protocol_dist"] = protocol_dist
            if protocol_dist:
                labels = [p["protocol"] for p in protocol_dist[:10]]
                values = [p["total_bytes"] for p in protocol_dist[:10]]
                png_bytes = render_bar_chart(
                    labels, values,
                    title="Protocol Distribution by Traffic Volume",
                    xlabel="Bytes",
                )
                context["charts"]["protocol_dist"] = base64.b64encode(png_bytes).decode()
        except Exception as exc:
            logger.error("R-01 protocol_dist fetch failed: %s", exc, exc_info=True)

        # R-01 Enhanced: Per-site traffic flow summary
        for site in DEFAULT_SITES:
            try:
                from app.opensearch import traffic_flow as tf_qb
                site_data = await tf_qb.flow_summary(
                    gte_ms=gte_ms, lte_ms=lte_ms, site_name=site
                )
                context["per_site_traffic"][site] = site_data
            except Exception as exc:
                logger.error("R-01 per-site traffic fetch failed for %s: %s", site, exc, exc_info=True)

    # ── R-02: Resource Usage ─────────────────────────────────────
    if report_type in ("R-02", "R-08"):
        try:
            from app.opensearch import ha as ha_qb

            devices = await ha_qb.current_device_status(gte_ms=gte_ms, lte_ms=lte_ms)
            context["devices"] = devices

            timeline = await ha_qb.resource_timeline(
                gte_ms=gte_ms, lte_ms=lte_ms, interval="15m"
            )
            if timeline["cpu"]:
                png_bytes = render_timeseries_chart(
                    timeline["cpu"],
                    title="CPU Usage Over Time",
                    ylabel="CPU %",
                    series_key="device",
                )
                context["charts"]["cpu_timeline"] = base64.b64encode(png_bytes).decode()
        except Exception as exc:
            logger.error("R-02 resource usage fetch failed: %s", exc, exc_info=True)

    # ── R-03: VPN Users ──────────────────────────────────────────
    if report_type in ("R-03", "R-08"):
        try:
            from app.opensearch import sslvpn as sslvpn_qb
            from app.opensearch import ipsec as ipsec_qb

            ssl_count = await sslvpn_qb.all_sslvpn_users_count(
                gte_ms=gte_ms, lte_ms=lte_ms, site_names=settings.sslvpn_sites_list
            )
            ipsec_count = await ipsec_qb.active_ipsec_users_count(
                gte_ms=gte_ms, lte_ms=lte_ms
            )
            context["ssl_vpn_count"] = ssl_count
            context["ipsec_vpn_count"] = ipsec_count

            # VPN bar chart
            if ssl_count > 0 or ipsec_count > 0:
                png_bytes = render_vpn_bar_chart(
                    ssl_count=ssl_count or 0,
                    ipsec_count=ipsec_count or 0,
                )
                context["charts"]["vpn_users"] = base64.b64encode(png_bytes).decode()
        except Exception as exc:
            logger.error("R-03 VPN users fetch failed: %s", exc, exc_info=True)

    # ── R-04: SD-WAN SLA ────────────────────────────────────────
    if report_type in ("R-04", "R-08"):
        try:
            from app.opensearch import sdwan as sdwan_qb
            context["sla_data"] = {}
            context["sla_charts"] = {}

            sla_sites = ["Site_FGT-DC", "Site_FGT-DRC", "Site_FGT_Office"]

            for site in sla_sites:
                site_sla = {}

                # SLA Timeline for each metric
                for metric in ("latency", "jitter", "packet_loss"):
                    try:
                        timeline_data = await sdwan_qb.sla_timeline(
                            gte_ms=gte_ms, lte_ms=lte_ms,
                            site_name=site, metric=metric, interval="5m",
                        )
                        site_sla[f"{metric}_timeline"] = timeline_data

                        # Render multi-line timeseries chart for latency and jitter
                        if timeline_data and metric in ("latency", "jitter"):
                            chart_key = f"sla_{metric}_{site}"
                            png_bytes = render_timeseries_chart(
                                timeline_data,
                                title=f"SD-WAN {metric.title()} — {site}",
                                ylabel=metric.title(),
                                series_key="label",
                            )
                            context["sla_charts"][chart_key] = base64.b64encode(png_bytes).decode()
                    except Exception as exc:
                        logger.error("R-04 SLA %s timeline fetch failed for %s: %s", metric, site, exc, exc_info=True)

                # SLA Summary (compliance data)
                try:
                    summary = await sdwan_qb.sla_summary(
                        gte_ms=gte_ms, lte_ms=lte_ms, site_name=site,
                    )
                    site_sla["summary"] = summary
                except Exception as exc:
                    logger.error("R-04 SLA summary fetch failed for %s: %s", site, exc, exc_info=True)

                context["sla_data"][site] = site_sla

        except Exception as exc:
            logger.error("R-04 SD-WAN SLA data fetch failed: %s", exc, exc_info=True)

    # ── R-05: Traffic Inbound ────────────────────────────────────
    if report_type in ("R-05", "R-08"):
        try:
            from app.opensearch import traffic_inbound as ti_qb
            context["inbound_data"] = {}

            inbound_sites = ["Site_FGT-DC", "Site_FGT-DRC"]
            for site in inbound_sites:
                try:
                    site_summary = await ti_qb.flow_summary(
                        gte_ms=gte_ms, lte_ms=lte_ms, site_name=site,
                    )
                    context["inbound_data"][site] = site_summary

                    # Render bar charts for inbound data
                    # Top Services
                    if site_summary.get("top_services"):
                        svc_labels = [s["service_name"] for s in site_summary["top_services"][:10]]
                        svc_values = [s["total_bytes"] for s in site_summary["top_services"][:10]]
                        png_bytes = render_bar_chart(
                            svc_labels, svc_values,
                            title=f"Top Inbound Services — {site}",
                            xlabel="Bytes",
                        )
                        context["charts"][f"inbound_services_{site}"] = base64.b64encode(png_bytes).decode()

                    # Top Source AS Orgs
                    if site_summary.get("top_src_as_org"):
                        org_labels = [o["org_name"] for o in site_summary["top_src_as_org"][:10]]
                        org_values = [o["total_bytes"] for o in site_summary["top_src_as_org"][:10]]
                        png_bytes = render_bar_chart(
                            org_labels, org_values,
                            title=f"Top Source AS Organizations — {site}",
                            xlabel="Bytes",
                        )
                        context["charts"][f"inbound_src_org_{site}"] = base64.b64encode(png_bytes).decode()

                    # Top Source Countries
                    if site_summary.get("top_src_as_country"):
                        country_labels = [c["country"] for c in site_summary["top_src_as_country"][:10]]
                        country_values = [c["total_bytes"] for c in site_summary["top_src_as_country"][:10]]
                        png_bytes = render_bar_chart(
                            country_labels, country_values,
                            title=f"Top Source Countries — {site}",
                            xlabel="Bytes",
                        )
                        context["charts"][f"inbound_country_{site}"] = base64.b64encode(png_bytes).decode()

                    # Protocol Distribution
                    if site_summary.get("protocol_dist"):
                        proto_labels = [p["protocol"] for p in site_summary["protocol_dist"][:10]]
                        proto_values = [p["total_bytes"] for p in site_summary["protocol_dist"][:10]]
                        png_bytes = render_bar_chart(
                            proto_labels, proto_values,
                            title=f"Inbound Protocol Distribution — {site}",
                            xlabel="Bytes",
                        )
                        context["charts"][f"inbound_protocol_{site}"] = base64.b64encode(png_bytes).decode()

                    # Egress Breakdown
                    if site_summary.get("egress_breakdown"):
                        egr_labels = [e["interface"] for e in site_summary["egress_breakdown"][:10]]
                        egr_values = [e["total_bytes"] for e in site_summary["egress_breakdown"][:10]]
                        png_bytes = render_bar_chart(
                            egr_labels, egr_values,
                            title=f"Inbound Egress Interface Breakdown — {site}",
                            xlabel="Bytes",
                        )
                        context["charts"][f"inbound_egress_{site}"] = base64.b64encode(png_bytes).decode()

                except Exception as exc:
                    logger.error("R-05 inbound traffic fetch failed for %s: %s", site, exc, exc_info=True)

        except Exception as exc:
            logger.error("R-05 traffic inbound data fetch failed: %s", exc, exc_info=True)

    # ── R-06: Traffic Internal ───────────────────────────────────
    if report_type in ("R-06", "R-08"):
        try:
            from app.opensearch import traffic_internal as tint_qb
            context["internal_data"] = {}

            internal_sites = ["Site_FGT-DC", "Site_FGT-DRC", "Site_FGT_Office"]
            for site in internal_sites:
                try:
                    site_summary = await tint_qb.flow_summary(
                        gte_ms=gte_ms, lte_ms=lte_ms, site_name=site,
                    )
                    context["internal_data"][site] = site_summary
                except Exception as exc:
                    logger.error("R-06 internal traffic fetch failed for %s: %s", site, exc, exc_info=True)

        except Exception as exc:
            logger.error("R-06 traffic internal data fetch failed: %s", exc, exc_info=True)

    # ── R-07: Executive Summary (combines key metrics) ──────────
    if report_type in ("R-07", "R-08"):
        try:
            context["executive_summary"] = _build_executive_summary(context)
        except Exception as exc:
            logger.error("R-07 executive summary build failed: %s", exc, exc_info=True)

    return context


def _build_executive_summary(context: dict) -> dict:
    """
    Build a single-page KPI summary from data already gathered in context.
    Works for R-07 standalone or R-08 (where context has all report data).
    """
    summary: dict = {}

    # Top 5 apps from R-01 data
    top_apps = context.get("top_applications", [])
    if top_apps:
        summary["top_5_apps"] = [
            {"app": a.get("application", ""), "bytes": a.get("total_bytes", 0)}
            for a in top_apps[:5]
        ]

    # Device health from R-02 data
    devices = context.get("devices", [])
    if devices:
        healthy = sum(1 for d in devices if d.get("cpu_usage", 0) < 80 and d.get("mem_usage", 0) < 80)
        summary["device_health"] = {
            "total_devices": len(devices),
            "healthy": healthy,
            "degraded": len(devices) - healthy,
        }

    # VPN counts from R-03 data
    summary["vpn_summary"] = {
        "ssl_vpn_count": context.get("ssl_vpn_count", 0),
        "ipsec_vpn_count": context.get("ipsec_vpn_count", 0),
    }

    # SLA summary from R-04 data
    sla_data = context.get("sla_data", {})
    if sla_data:
        sla_summaries = {}
        for site, site_sla in sla_data.items():
            site_summary = site_sla.get("summary", {})
            if site_summary:
                sla_summaries[site] = {
                    "avg_latency": site_summary.get("avg_latency", []),
                    "avg_jitter": site_summary.get("avg_jitter", []),
                    "avg_packet_loss": site_summary.get("avg_packet_loss", []),
                    "labels": site_summary.get("labels", []),
                }
        summary["sla_summary"] = sla_summaries

    # Total throughput
    summary["total_throughput_bytes"] = context.get("total_throughput_bytes", 0)

    return summary


# ─────────────────────────────────────────────────────────────────
# Format Generators
# ─────────────────────────────────────────────────────────────────


def _generate_html(context: dict, output_path: Path) -> None:
    """Generate a self-contained HTML report using Jinja2."""
    from jinja2 import Template

    template_str = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ report_title }} — NOD</title>
<style>
    /* ── Design Tokens ─────────────────────────────────────── */
    :root {
        --nod-primary: #2563eb;
        --nod-primary-light: #dbeafe;
        --nod-primary-dark: #1d4ed8;
        --nod-success: #10b981;
        --nod-success-light: #d1fae5;
        --nod-warning: #f59e0b;
        --nod-warning-light: #fef3c7;
        --nod-danger: #ef4444;
        --nod-danger-light: #fee2e2;
        --nod-bg: #f8fafc;
        --nod-card: #ffffff;
        --nod-text: #1e293b;
        --nod-text-secondary: #64748b;
        --nod-border: #e2e8f0;
        --nod-border-radius: 8px;
        --nod-shadow: 0 1px 3px rgba(0,0,0,0.08), 0 1px 2px rgba(0,0,0,0.06);
        --nod-font: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
    }

    /* ── Base ───────────────────────────────────────────────── */
    *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
    body {
        font-family: var(--nod-font);
        background: var(--nod-bg);
        color: var(--nod-text);
        font-size: 14px;
        line-height: 1.6;
        margin: 0;
        padding: 0;
    }
    .report-wrapper { max-width: 1100px; margin: 0 auto; padding: 0 32px 48px; }

    /* ── Header ─────────────────────────────────────────────── */
    .nod-header {
        background: linear-gradient(135deg, var(--nod-primary) 0%, var(--nod-primary-dark) 100%);
        color: #ffffff;
        padding: 32px 40px;
        margin: 0 -32px 0;
    }
    .nod-header-inner { max-width: 1100px; margin: 0 auto; }
    .nod-header .logo-area {
        display: flex;
        align-items: center;
        gap: 16px;
        margin-bottom: 12px;
    }
    .nod-header .logo-icon {
        width: 48px; height: 48px;
        background: rgba(255,255,255,0.2);
        border-radius: 12px;
        display: flex; align-items: center; justify-content: center;
        font-size: 24px; font-weight: 700; letter-spacing: -1px;
    }
    .nod-header .brand-text {
        font-size: 13px; font-weight: 600;
        text-transform: uppercase; letter-spacing: 2px;
        opacity: 0.9;
    }
    .nod-header h1 {
        font-size: 26px; font-weight: 700; margin: 0 0 4px;
        border: none; padding: 0; color: #ffffff;
    }
    .nod-header .subtitle {
        font-size: 16px; font-weight: 400; opacity: 0.85;
    }

    /* ── Meta Bar ───────────────────────────────────────────── */
    .meta-bar {
        background: var(--nod-card);
        border: 1px solid var(--nod-border);
        border-radius: var(--nod-border-radius);
        padding: 14px 24px;
        margin: 24px 0;
        display: flex;
        flex-wrap: wrap;
        gap: 8px 32px;
        font-size: 13px;
        color: var(--nod-text-secondary);
    }
    .meta-bar .meta-item { display: flex; align-items: center; gap: 6px; }
    .meta-bar .meta-label { font-weight: 600; color: var(--nod-text); }

    /* ── Section Headings ───────────────────────────────────── */
    .section-title {
        font-size: 20px;
        font-weight: 700;
        color: var(--nod-text);
        margin: 40px 0 16px;
        padding-bottom: 8px;
        border-bottom: 3px solid var(--nod-primary);
        page-break-after: avoid;
    }
    .section-subtitle {
        font-size: 15px;
        font-weight: 600;
        color: var(--nod-text);
        margin: 24px 0 12px;
        page-break-after: avoid;
    }

    /* ── KPI Cards ──────────────────────────────────────────── */
    .kpi-grid {
        display: flex;
        gap: 16px;
        flex-wrap: wrap;
        margin: 20px 0 32px;
    }
    .kpi-card {
        flex: 1 1 160px;
        background: var(--nod-card);
        border: 1px solid var(--nod-border);
        border-radius: var(--nod-border-radius);
        padding: 20px;
        text-align: center;
        box-shadow: var(--nod-shadow);
    }
    .kpi-card .kpi-label {
        font-size: 11px;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        color: var(--nod-text-secondary);
        margin-bottom: 8px;
    }
    .kpi-card .kpi-value {
        font-size: 28px;
        font-weight: 700;
        color: var(--nod-primary);
        line-height: 1.2;
    }
    .kpi-card .kpi-value.success { color: var(--nod-success); }
    .kpi-card .kpi-value.warning { color: var(--nod-warning); }
    .kpi-card .kpi-value.danger { color: var(--nod-danger); }

    /* ── Charts Grid ────────────────────────────────────────── */
    .chart-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 20px;
        margin: 20px 0 32px;
    }
    .chart-card {
        background: var(--nod-card);
        border: 1px solid var(--nod-border);
        border-radius: var(--nod-border-radius);
        padding: 16px;
        box-shadow: var(--nod-shadow);
        page-break-inside: avoid;
    }
    .chart-card.full-width { grid-column: 1 / -1; }
    .chart-card h3 {
        font-size: 14px;
        font-weight: 600;
        color: var(--nod-text);
        margin-bottom: 12px;
        padding-bottom: 8px;
        border-bottom: 1px solid var(--nod-border);
    }
    .chart-card img {
        max-width: 100%;
        height: auto;
        display: block;
        border-radius: 4px;
    }

    /* ── Tables ─────────────────────────────────────────────── */
    .table-card {
        background: var(--nod-card);
        border: 1px solid var(--nod-border);
        border-radius: var(--nod-border-radius);
        box-shadow: var(--nod-shadow);
        margin: 20px 0 32px;
        overflow: hidden;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        font-size: 13px;
    }
    thead th {
        background: var(--nod-primary-light);
        color: var(--nod-primary-dark);
        font-weight: 600;
        text-align: left;
        padding: 12px 16px;
        border-bottom: 2px solid var(--nod-primary);
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.3px;
    }
    tbody td {
        padding: 10px 16px;
        border-bottom: 1px solid var(--nod-border);
    }
    tbody tr:last-child td { border-bottom: none; }
    tbody tr:hover { background: #f1f5f9; }

    /* Color-coded status cells */
    .status-green { color: var(--nod-success); font-weight: 600; }
    .status-amber { color: var(--nod-warning); font-weight: 600; }
    .status-red { color: var(--nod-danger); font-weight: 600; }

    /* ── Footer ─────────────────────────────────────────────── */
    .nod-footer {
        margin-top: 48px;
        padding: 24px 0;
        border-top: 2px solid var(--nod-border);
        text-align: center;
        font-size: 11px;
        color: var(--nod-text-secondary);
        page-break-before: auto;
    }
    .nod-footer .footer-brand {
        font-weight: 700;
        color: var(--nod-primary);
        font-size: 13px;
        margin-bottom: 6px;
    }
    .nod-footer .confidential {
        font-weight: 600;
        color: var(--nod-danger);
        margin-top: 8px;
        text-transform: uppercase;
        letter-spacing: 1px;
    }

    /* ── Per-site chart section ──────────────────────────────── */
    .site-section { margin-bottom: 32px; }
    .site-section .site-header {
        font-size: 16px;
        font-weight: 600;
        color: var(--nod-primary);
        padding: 8px 16px;
        background: var(--nod-primary-light);
        border-radius: var(--nod-border-radius) var(--nod-border-radius) 0 0;
        border: 1px solid var(--nod-border);
        border-bottom: none;
    }
    .site-section .site-body {
        background: var(--nod-card);
        border: 1px solid var(--nod-border);
        border-radius: 0 0 var(--nod-border-radius) var(--nod-border-radius);
        padding: 16px;
    }

    /* ── Print Styles ────────────────────────────────────────── */
    @page {
        size: A4;
        margin: 20mm 15mm 25mm 15mm;
        @bottom-center {
            content: "NOD Report — Page " counter(page) " of " counter(pages);
            font-family: var(--nod-font);
            font-size: 9px;
            color: #94a3b8;
        }
        @bottom-right {
            content: "CONFIDENTIAL";
            font-family: var(--nod-font);
            font-size: 8px;
            color: #ef4444;
            font-weight: 600;
        }
    }
    @page :first {
        margin-top: 15mm;
    }
    @media print {
        body { background: #fff; }
        .nod-header { margin: -20mm -15mm 0; padding: 24px 40px; }
        .report-wrapper { padding: 0; }
        .section-title { page-break-after: avoid; }
        .chart-card, .table-card { page-break-inside: avoid; }
        .page-break { page-break-before: always; }
    }
</style>
</head>
<body>

<!-- ═══════════════ HEADER ═══════════════ -->
<div class="nod-header">
    <div class="nod-header-inner">
        <div class="logo-area">
            <div class="logo-icon">NOD</div>
            <div class="brand-text">Network Observability Dashboard</div>
        </div>
        <h1>{{ report_title }}</h1>
        <div class="subtitle">Comprehensive Network Intelligence Report</div>
    </div>
</div>

<div class="report-wrapper">

<!-- ═══════════════ META BAR ═══════════════ -->
<div class="meta-bar">
    <div class="meta-item">
        <span class="meta-label">Generated:</span> {{ generated_at }}
    </div>
    <div class="meta-item">
        <span class="meta-label">Author:</span> {{ generated_by }}
    </div>
    <div class="meta-item">
        <span class="meta-label">Job ID:</span> {{ job_id }}
    </div>
    <div class="meta-item">
        <span class="meta-label">Time Range:</span> {{ time_range }}
    </div>
</div>

<!-- ═══════════════ EXECUTIVE SUMMARY ═══════════════ -->
{% if executive_summary %}
<h2 class="section-title">Executive Summary</h2>
<div class="kpi-grid">
    <div class="kpi-card">
        <div class="kpi-label">Total Throughput</div>
        <div class="kpi-value">{{ executive_summary.total_throughput_bytes | default(total_throughput_bytes | default(0)) }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">SSL VPN Users</div>
        <div class="kpi-value">{{ executive_summary.vpn_summary.ssl_vpn_count | default(ssl_vpn_count | default(0)) }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">IPsec VPN Users</div>
        <div class="kpi-value">{{ executive_summary.vpn_summary.ipsec_vpn_count | default(ipsec_vpn_count | default(0)) }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">Total Devices</div>
        <div class="kpi-value">{{ executive_summary.device_health.total_devices | default(0) }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">SLA Sites</div>
        <div class="kpi-value">{{ executive_summary.sla_summary | length | default(0) }}</div>
    </div>
</div>

{% if executive_summary.top_5_apps %}
<h3 class="section-subtitle">Top 5 Applications</h3>
<div class="table-card">
<table>
    <thead><tr><th>#</th><th>Application</th><th>Total Bytes</th></tr></thead>
    <tbody>
    {% for app in executive_summary.top_5_apps %}
    <tr>
        <td>{{ loop.index }}</td>
        <td>{{ app.app }}</td>
        <td>{{ app.bytes }}</td>
    </tr>
    {% endfor %}
    </tbody>
</table>
</div>
{% endif %}

{% if executive_summary.device_health %}
<h3 class="section-subtitle">Device Health Overview</h3>
<div class="kpi-grid">
    <div class="kpi-card">
        <div class="kpi-label">Total Devices</div>
        <div class="kpi-value">{{ executive_summary.device_health.total_devices }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">Healthy</div>
        <div class="kpi-value success">{{ executive_summary.device_health.healthy }}</div>
    </div>
    <div class="kpi-card">
        <div class="kpi-label">Degraded</div>
        <div class="kpi-value warning">{{ executive_summary.device_health.degraded }}</div>
    </div>
</div>
{% endif %}

{% if executive_summary.sla_summary %}
<h3 class="section-subtitle">SLA Summary</h3>
<div class="table-card">
<table>
    <thead><tr><th>Site</th><th>Avg Latency (ms)</th><th>Avg Jitter (ms)</th><th>Avg Packet Loss (%)</th></tr></thead>
    <tbody>
    {% for site, sla in executive_summary.sla_summary.items() %}
    <tr>
        <td><strong>{{ site }}</strong></td>
        <td>{{ sla.avg_latency }}</td>
        <td>{{ sla.avg_jitter }}</td>
        <td>{{ sla.avg_packet_loss }}</td>
    </tr>
    {% endfor %}
    </tbody>
</table>
</div>
{% endif %}
{% endif %}

<!-- ═══════════════ R-01: TRAFFIC FLOW ═══════════════ -->
{% if charts.top_applications or charts.throughput or charts.top_as_orgs or charts.top_countries or charts.protocol_dist %}
<div class="page-break"></div>
<h2 class="section-title">Traffic Flow Analysis</h2>

<div class="chart-grid">
    {% if charts.top_applications %}
    <div class="chart-card">
        <h3>Top 10 Applications by Traffic Volume</h3>
        <img src="data:image/png;base64,{{ charts.top_applications }}" alt="Top Applications Chart">
    </div>
    {% endif %}

    {% if charts.throughput %}
    <div class="chart-card">
        <h3>Throughput Over Time</h3>
        <img src="data:image/png;base64,{{ charts.throughput }}" alt="Throughput Chart">
    </div>
    {% endif %}

    {% if charts.top_as_orgs %}
    <div class="chart-card">
        <h3>Top 10 Destination AS Organizations</h3>
        <img src="data:image/png;base64,{{ charts.top_as_orgs }}" alt="Top AS Orgs Chart">
    </div>
    {% endif %}

    {% if charts.top_countries %}
    <div class="chart-card">
        <h3>Top 10 Destination Countries</h3>
        <img src="data:image/png;base64,{{ charts.top_countries }}" alt="Top Countries Chart">
    </div>
    {% endif %}

    {% if charts.protocol_dist %}
    <div class="chart-card full-width">
        <h3>Protocol Distribution by Traffic Volume</h3>
        <img src="data:image/png;base64,{{ charts.protocol_dist }}" alt="Protocol Distribution Chart">
    </div>
    {% endif %}
</div>
{% endif %}

<!-- ═══════════════ PER-SITE TRAFFIC TABLE ═══════════════ -->
{% if per_site_traffic %}
<h2 class="section-title">Per-Site Traffic Summary</h2>
<div class="table-card">
<table>
    <thead>
    <tr><th>Site</th><th>Total Bytes</th><th>Top Application</th><th>Sessions</th></tr>
    </thead>
    <tbody>
    {% for site, site_data in per_site_traffic.items() %}
    <tr>
        <td><strong>{{ site }}</strong></td>
        <td>{{ site_data.total_bytes | default('N/A') }}</td>
        <td>{{ site_data.top_apps[0].app_name | default('N/A') }}</td>
        <td>{{ site_data.total_sessions | default('N/A') }}</td>
    </tr>
    {% endfor %}
    </tbody>
</table>
</div>
{% endif %}

<!-- ═══════════════ R-02: DEVICE STATUS ═══════════════ -->
{% if devices %}
<div class="page-break"></div>
<h2 class="section-title">Device Resource Status</h2>
<div class="table-card">
<table>
    <thead>
    <tr><th>Device</th><th>Hostname</th><th>CPU %</th><th>Memory %</th><th>Sessions</th><th>Sync Status</th></tr>
    </thead>
    <tbody>
    {% for d in devices %}
    <tr>
        <td><strong>{{ d.device }}</strong></td>
        <td>{{ d.hostname }}</td>
        <td class="{% if d.cpu_usage < 60 %}status-green{% elif d.cpu_usage < 80 %}status-amber{% else %}status-red{% endif %}">{{ d.cpu_usage }}</td>
        <td class="{% if d.mem_usage < 60 %}status-green{% elif d.mem_usage < 80 %}status-amber{% else %}status-red{% endif %}">{{ d.mem_usage }}</td>
        <td>{{ d.session_count }}</td>
        <td>{{ d.sync_status }}</td>
    </tr>
    {% endfor %}
    </tbody>
</table>
</div>

{% if charts.cpu_timeline %}
<div class="chart-grid">
    <div class="chart-card full-width">
        <h3>CPU Usage Over Time</h3>
        <img src="data:image/png;base64,{{ charts.cpu_timeline }}" alt="CPU Timeline Chart">
    </div>
</div>
{% endif %}
{% endif %}

<!-- ═══════════════ R-04: SD-WAN SLA ═══════════════ -->
{% if sla_data or sla_charts %}
<div class="page-break"></div>
<h2 class="section-title">SD-WAN SLA Monitoring</h2>

{% if sla_charts %}
<div class="chart-grid">
    {% for chart_key, chart_val in sla_charts.items() %}
    <div class="chart-card">
        <h3>{{ chart_key | replace('_', ' ') | title }}</h3>
        <img src="data:image/png;base64,{{ chart_val }}" alt="{{ chart_key }}">
    </div>
    {% endfor %}
</div>
{% endif %}

{% if sla_data %}
<h3 class="section-subtitle">SLA Compliance Summary</h3>
<div class="table-card">
<table>
    <thead>
    <tr><th>Site</th><th>Avg Latency (ms)</th><th>Avg Jitter (ms)</th><th>Avg Packet Loss (%)</th></tr>
    </thead>
    <tbody>
    {% for site, site_sla in sla_data.items() %}
    {% if site_sla.summary %}
    <tr>
        <td><strong>{{ site }}</strong></td>
        <td>{{ site_sla.summary.avg_latency | default('N/A') }}</td>
        <td>{{ site_sla.summary.avg_jitter | default('N/A') }}</td>
        <td>{{ site_sla.summary.avg_packet_loss | default('N/A') }}</td>
    </tr>
    {% endif %}
    {% endfor %}
    </tbody>
</table>
</div>
{% endif %}
{% endif %}

<!-- ═══════════════ R-05: TRAFFIC INBOUND ═══════════════ -->
{% if inbound_data %}
<div class="page-break"></div>
<h2 class="section-title">Traffic Inbound Analysis</h2>

{% for chart_key, chart_val in charts.items() if chart_key.startswith('inbound_') %}
{% if loop.index0 % 3 == 0 %}<div class="chart-grid">{% endif %}
    <div class="chart-card">
        <h3>{{ chart_key | replace('_', ' ') | title }}</h3>
        <img src="data:image/png;base64,{{ chart_val }}" alt="{{ chart_key }}">
    </div>
{% if loop.index0 % 3 == 2 or loop.last %}</div>{% endif %}
{% endfor %}

{% for site, site_data in inbound_data.items() %}
{% if site_data.top_services %}
<div class="site-section">
    <div class="site-header">{{ site }} — Top Inbound Services</div>
    <div class="site-body">
        <div class="table-card" style="margin:0;box-shadow:none;">
        <table>
            <thead><tr><th>#</th><th>Service</th><th>Total Bytes</th></tr></thead>
            <tbody>
            {% for svc in site_data.top_services[:10] %}
            <tr><td>{{ loop.index }}</td><td>{{ svc.service_name }}</td><td>{{ svc.total_bytes }}</td></tr>
            {% endfor %}
            </tbody>
        </table>
        </div>
    </div>
</div>
{% endif %}
{% endfor %}
{% endif %}

<!-- ═══════════════ R-06: TRAFFIC INTERNAL ═══════════════ -->
{% if internal_data %}
<div class="page-break"></div>
<h2 class="section-title">Traffic Internal Analysis</h2>

{% for site, site_data in internal_data.items() %}
{% if site_data.top_services %}
<div class="site-section">
    <div class="site-header">{{ site }} — Top Internal Services</div>
    <div class="site-body">
        <div class="table-card" style="margin:0;box-shadow:none;">
        <table>
            <thead><tr><th>#</th><th>Service</th><th>Total Bytes</th></tr></thead>
            <tbody>
            {% for svc in site_data.top_services[:10] %}
            <tr><td>{{ loop.index }}</td><td>{{ svc.service_name }}</td><td>{{ svc.total_bytes }}</td></tr>
            {% endfor %}
            </tbody>
        </table>
        </div>
    </div>
</div>
{% endif %}
{% endfor %}
{% endif %}

<!-- ═══════════════ R-03: VPN USERS ═══════════════ -->
{% if charts.vpn_users %}
<div class="page-break"></div>
<h2 class="section-title">Active VPN Users</h2>
<div class="chart-grid">
    <div class="chart-card full-width">
        <h3>VPN Users Overview</h3>
        <img src="data:image/png;base64,{{ charts.vpn_users }}" alt="VPN Users Chart">
    </div>
</div>
{% endif %}

<!-- ═══════════════ FOOTER ═══════════════ -->
<div class="nod-footer">
    <div class="footer-brand">NOD — Network Observability Dashboard</div>
    <div>Report Type: {{ report_title }} | Job ID: {{ job_id }}</div>
    <div>Generated {{ generated_at }} by {{ generated_by }}</div>
    <div class="confidential">Document Classification: Internal — Confidential</div>
</div>

</div><!-- /report-wrapper -->
</body>
</html>"""

    template = Template(template_str)
    html_content = template.render(**context)
    output_path.write_text(html_content, encoding="utf-8")


def _generate_pdf(context: dict, output_path: Path) -> None:
    """Generate PDF via WeasyPrint from HTML template."""
    import tempfile

    # Generate HTML first, then convert to PDF
    html_path = output_path.with_suffix(".html.tmp")
    _generate_html(context, html_path)

    try:
        HTML(filename=str(html_path)).write_pdf(str(output_path))
    finally:
        if html_path.exists():
            html_path.unlink()  # Clean up temp HTML


def _generate_docx(context: dict, output_path: Path) -> None:
    """Generate DOCX report using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(context.get("report_title", "NOD Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    doc.add_paragraph(
        f"Generated: {context.get('generated_at', '')} | By: {context.get('generated_by', '')} | ID: {context.get('job_id', '')}"
    ).style = "Normal"
    doc.add_paragraph(f"Time Range: {context.get('time_range', '')}").style = "Normal"

    # ── R-07 / R-08: Executive Summary ──
    doc.add_heading("R-07: Executive Summary", level=1)
    doc.add_paragraph(f"Total Throughput: {context.get('total_throughput_bytes', 0):,} bytes")
    doc.add_paragraph(f"Active SSL VPN Users: {context.get('ssl_vpn_count', 0)}")
    doc.add_paragraph(f"Active IPsec VPN Users: {context.get('ipsec_vpn_count', 0)}")

    # R-07: Executive Summary KPIs
    exec_summary = context.get("executive_summary", {})
    if exec_summary:
        doc.add_heading("Executive Summary KPIs", level=1)
        if exec_summary.get("top_5_apps"):
            doc.add_heading("Top 5 Applications", level=2)
            table = doc.add_table(rows=1, cols=2, style="Table Grid")
            table.rows[0].cells[0].text = "Application"
            table.rows[0].cells[1].text = "Total Bytes"
            for app in exec_summary["top_5_apps"]:
                row = table.add_row()
                row.cells[0].text = app.get("app", "")
                row.cells[1].text = str(app.get("bytes", 0))

        if exec_summary.get("device_health"):
            dh = exec_summary["device_health"]
            doc.add_heading("Device Health", level=2)
            doc.add_paragraph(
                f"Total: {dh.get('total_devices', 0)} | Healthy: {dh.get('healthy', 0)} | Degraded: {dh.get('degraded', 0)}"
            )

        if exec_summary.get("vpn_summary"):
            vs = exec_summary["vpn_summary"]
            doc.add_heading("VPN Summary", level=2)
            doc.add_paragraph(f"SSL VPN: {vs.get('ssl_vpn_count', 0)} | IPsec VPN: {vs.get('ipsec_vpn_count', 0)}")

    # ── R-01: Per-site traffic ──
    per_site = context.get("per_site_traffic", {})
    if per_site:
        doc.add_heading("R-01: Per-Site Traffic Summary", level=1)
        for site, site_data in per_site.items():
            doc.add_heading(site, level=2)
            top_apps = site_data.get("top_apps", [])
            if top_apps:
                table = doc.add_table(rows=1, cols=3, style="Table Grid")
                table.rows[0].cells[0].text = "Application"
                table.rows[0].cells[1].text = "Bytes"
                table.rows[0].cells[2].text = "Mbps"
                for a in top_apps[:5]:
                    row = table.add_row()
                    row.cells[0].text = a.get("app_name", "")
                    row.cells[1].text = str(a.get("total_bytes", 0))
                    row.cells[2].text = f"{a.get('speed_mbps', 0):.2f}"

    # ── R-04: SD-WAN SLA Data ──
    sla_data = context.get("sla_data", {})
    if sla_data:
        doc.add_heading("R-04: SD-WAN SLA Data", level=1)
        for site, site_sla in sla_data.items():
            doc.add_heading(site, level=2)
            summary = site_sla.get("summary", {})
            if summary:
                doc.add_paragraph(f"Labels: {', '.join(summary.get('labels', []))}")
                doc.add_paragraph(f"Avg Latency: {summary.get('avg_latency', [])}")
                doc.add_paragraph(f"Avg Jitter: {summary.get('avg_jitter', [])}")
                doc.add_paragraph(f"Avg Packet Loss: {summary.get('avg_packet_loss', [])}")

    # ── R-05: Traffic Inbound Data ──
    inbound = context.get("inbound_data", {})
    if inbound:
        doc.add_heading("R-05: Traffic Inbound Data", level=1)
        for site, site_data in inbound.items():
            doc.add_heading(site, level=2)
            top_svcs = site_data.get("top_services", [])
            if top_svcs:
                table = doc.add_table(rows=1, cols=2, style="Table Grid")
                table.rows[0].cells[0].text = "Service"
                table.rows[0].cells[1].text = "Bytes"
                for s in top_svcs[:5]:
                    row = table.add_row()
                    row.cells[0].text = s.get("service_name", "")
                    row.cells[1].text = str(s.get("total_bytes", 0))

    # ── R-06: Traffic Internal Data ──
    internal = context.get("internal_data", {})
    if internal:
        doc.add_heading("R-06: Traffic Internal Data", level=1)
        for site, site_data in internal.items():
            doc.add_heading(site, level=2)
            top_svcs = site_data.get("top_services", [])
            if top_svcs:
                table = doc.add_table(rows=1, cols=2, style="Table Grid")
                table.rows[0].cells[0].text = "Service"
                table.rows[0].cells[1].text = "Bytes"
                for s in top_svcs[:5]:
                    row = table.add_row()
                    row.cells[0].text = s.get("service_name", "")
                    row.cells[1].text = str(s.get("total_bytes", 0))

    # ── R-02: Device Resource Status ──
    devices = context.get("devices", [])
    if devices:
        doc.add_heading("R-02: Device Resource Status", level=1)
        table = doc.add_table(rows=1, cols=5, style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Device"
        hdr_cells[1].text = "CPU %"
        hdr_cells[2].text = "Memory %"
        hdr_cells[3].text = "Sessions"
        hdr_cells[4].text = "Sync Status"

        for d in devices:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{d.get('device', '')} ({d.get('hostname', '')})"
            row_cells[1].text = str(d.get("cpu_usage", ""))
            row_cells[2].text = str(d.get("mem_usage", ""))
            row_cells[3].text = str(d.get("session_count", ""))
            row_cells[4].text = d.get("sync_status", "")

    # ── R-03: VPN Users ──
    ssl_vpn = context.get("ssl_vpn_count", 0)
    ipsec_vpn = context.get("ipsec_vpn_count", 0)
    if ssl_vpn or ipsec_vpn:
        doc.add_heading("R-03: Active VPN Users", level=1)
        doc.add_paragraph(f"Active SSL VPN Users: {ssl_vpn}")
        doc.add_paragraph(f"Active IPsec VPN Users: {ipsec_vpn}")

    # ── R-08: Document Classification (All-in-One) ──
    doc.add_page_break()
    doc.add_heading("R-08: Document Classification", level=1)
    doc.add_paragraph("NOD — Network Observability Dashboard", style="Normal")
    doc.add_paragraph(
        "Document Classification: Internal — Confidential",
        style="Normal",
    )
    doc.add_paragraph(
        f"Report Type: {context.get('report_title', 'NOD Report')} | Job ID: {context.get('job_id', '')}",
        style="Normal",
    )
    doc.add_paragraph(
        f"Generated {context.get('generated_at', '')} by {context.get('generated_by', '')}",
        style="Normal",
    )

    # Legacy footer
    doc.add_page_break()
    doc.add_paragraph(
        "NOD — Network Observability Dashboard | Internal — Confidential",
        style="Normal",
    )

    doc.save(str(output_path))

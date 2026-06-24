#!/usr/bin/env python3
"""
DOCX Report Generator for NOD.

Takes the same context dict as the PDF/HTML generator and produces
a well-formatted Word document (.docx) using python-docx.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    logger.error("python-docx is not installed. Install it with: pip install python-docx")
    raise


# ── Colour palette ──────────────────────────────────────────────────────
PRIMARY_HEX = "2563EB"      # hsl(221.2 83.2% 53.3%)
MUTED_HEX   = "64748B"      # hsl(var(--muted-foreground))
SUCCESS_HEX = "16A34A"      # hsl(160.1 84.1% 39.4%)
DANGER_HEX  = "DC2626"      # hsl(0 72.2% 50.6%)
WARN_HEX    = "D97706"      # hsl(38.9 92.7% 50.2%)
WHITE_HEX   = "FFFFFF"
DARK_HEX    = "171717"      # hsl(0 0% 9%)


def _hex_color(hex_str: str) -> RGBColor:
    """Convert a hex colour string to python-docx RGBColor."""
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _add_styled_paragraph(doc, text: str, style_name: str = "Normal",
                           bold: bool = False, size: int = 10,
                           color_hex: str = DARK_HEX,
                           alignment: int = None, space_after: int = 4):
    """Add a paragraph with explicit formatting."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = _hex_color(color_hex)
    run.font.name = "Calibri"
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _add_heading_styled(doc, text: str, level: int = 1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _hex_color(DARK_HEX)
        run.font.name = "Calibri"
    return h


def _add_meta_table(doc, context: dict[str, Any]) -> None:
    """Add a 2-column metadata table (From, To, Generated, By, Job)."""
    rows_data = [
        ("From", context.get("time_range_start", "—")),
        ("To", context.get("time_range_end", "—")),
        ("Generated", context.get("generated_at", "—")),
        ("By", context.get("generated_by", "—")[:12] + "…"
             if len(context.get("generated_by", "")) > 12
             else context.get("generated_by", "—")),
        ("Job", context.get("job_id", "—")[:8] + "…"
             if len(context.get("job_id", "")) > 8
             else context.get("job_id", "—")),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows_data):
        cell_l = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_l.text = label
        cell_v.text = value
        for cell in (cell_l, cell_v):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
        # Bold the label
        for run in cell_l.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = _hex_color(MUTED_HEX)
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(4.5)


def _add_kpi(doc, label: str, value: str, note: str = "") -> None:
    """Add a single KPI card as a formatted paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = _hex_color(MUTED_HEX)
    run_label.font.name = "Calibri"
    run_val = p.add_run(value)
    run_val.bold = True
    run_val.font.size = Pt(14)
    run_val.font.color.rgb = _hex_color(PRIMARY_HEX)
    run_val.font.name = "Calibri"
    if note:
        p.add_run(f" — {note}").font.size = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)


def _add_data_table(doc, headers: list[str], rows: list[list[str]],
                    col_widths: list[float] | None = None) -> None:
    """Add a formatted table with header row."""
    if not rows:
        doc.add_paragraph("No data available.").paragraph_format.space_after = Pt(6)
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        _set_cell_shading(cell, PRIMARY_HEX)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = _hex_color(WHITE_HEX)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
            # Alternating row shading
            if i % 2 == 1:
                _set_cell_shading(cell, "F1F5F9")
    # Column widths
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Inches(w)


def _add_hbar_items(doc, items: list[dict] | None, limit: int = 10) -> None:
    """Add horizontal bar chart items as a bulleted list with values."""
    if not items:
        doc.add_paragraph("No data available for this timeframe.").paragraph_format.space_after = Pt(6)
        return
    for item in items[:limit]:
        label = item.get("label", "—")
        value = item.get("value", 0)
        if isinstance(value, (int, float)):
            # Format bytes nicely
            if value >= 1_000_000_000:
                val_str = f"{value / 1_000_000_000:.2f} GB"
            elif value >= 1_000_000:
                val_str = f"{value / 1_000_000:.2f} MB"
            elif value >= 1_000:
                val_str = f"{value / 1_000:.2f} KB"
            else:
                val_str = f"{value} B"
        else:
            val_str = str(value)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(f"{label}: {val_str}")
        run.font.size = Pt(9)
        run.font.name = "Calibri"


def generate_docx_report(context: dict[str, Any], output_path: Path) -> Path:
    """
    Generate a .docx Word document from the report context.
    Returns the output path.
    """
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # ── Title ──────────────────────────────────────────────────────
    title = context.get("report_title", "NOD Report")
    _add_heading_styled(doc, f"NOD — {title}", level=1)

    # ── Time range ────────────────────────────────────────────────
    _add_meta_table(doc, context)

    # ── Sections ──────────────────────────────────────────────────
    report_data: dict = context.get("report_data", {})
    rd = report_data

    # ── R-01 / R-08: Traffic Overview ─────────────────────────────
    to = rd.get("traffic_overview", {})
    if to:
        _add_heading_styled(doc, "Traffic Flow Analysis", level=2)

        if to.get("total_throughput_bytes") is not None:
            total = to["total_throughput_bytes"]
            if total >= 1_000_000_000:
                total_str = f"{total / 1_000_000_000:.2f} GB"
            elif total >= 1_000_000:
                total_str = f"{total / 1_000_000:.2f} MB"
            else:
                total_str = f"{total} B"
            _add_kpi(doc, "Total Throughput", total_str, "All sites combined")

        if to.get("top_applications"):
            _add_heading_styled(doc, "Top 10 Applications by Traffic Volume", level=3)
            _add_hbar_items(doc, to["top_applications"])

        if to.get("top_as_orgs"):
            _add_heading_styled(doc, "Top 10 Destination AS Organizations", level=3)
            _add_hbar_items(doc, to["top_as_orgs"])

        if to.get("top_countries"):
            _add_heading_styled(doc, "Top 10 Destination Countries", level=3)
            _add_hbar_items(doc, to["top_countries"])

        if to.get("per_site_summary"):
            _add_heading_styled(doc, "Per-Site Traffic Summary", level=3)
            headers = ["Site", "Total Traffic", "Throughput (Mbps)", "Sessions"]
            rows = []
            for ps in to["per_site_summary"]:
                tb = ps.get("total_bytes", 0)
                tb_str = f"{tb / 1_000_000_000:.2f} GB" if tb >= 1_000_000_000 else str(tb)
                mbps = ps.get("total_mbps", 0)
                rows.append([
                    ps.get("site", "—"),
                    tb_str,
                    f"{mbps:.2f} Mbps",
                    str(ps.get("sessions", "—")),
                ])
            _add_data_table(doc, headers, rows, col_widths=[1.2, 1.2, 1.2, 1.0])

    # ── R-02 / R-08: Resource Usage ───────────────────────────────
    ru = rd.get("resource_usage", {})
    if ru:
        _add_heading_styled(doc, "Resource Usage", level=2)
        if ru.get("cpu_usage_pct") is not None:
            _add_kpi(doc, "CPU", f"{ru['cpu_usage_pct']}%")
        if ru.get("memory_usage_pct") is not None:
            _add_kpi(doc, "Memory", f"{ru['memory_usage_pct']}%")
        if ru.get("disk_usage_pct") is not None:
            _add_kpi(doc, "Disk", f"{ru['disk_usage_pct']}%")

    # ── R-03 / R-08: VPN Users ────────────────────────────────────
    vu = rd.get("vpn_users", {})
    if vu:
        _add_heading_styled(doc, "Active VPN Users", level=2)
        if vu.get("total_count") is not None:
            _add_kpi(doc, "Total VPN Users", str(vu["total_count"]))
        if vu.get("ssl_count") is not None:
            _add_kpi(doc, "SSL VPN", str(vu["ssl_count"]))
        if vu.get("ipsec_count") is not None:
            _add_kpi(doc, "IPsec VPN", str(vu["ipsec_count"]))

    # ── R-04 / R-08: SD-WAN SLA ───────────────────────────────────
    sla = rd.get("sdwan_sla", {})
    if sla:
        _add_heading_styled(doc, "SD-WAN SLA", level=2)
        sla_links = sla.get("links", [])
        if sla_links:
            headers = ["Link", "Packet Loss", "Latency", "Status"]
            rows = []
            for link in sla_links:
                rows.append([
                    link.get("name", "—"),
                    f"{link.get('packet_loss', 0):.2f}%" if isinstance(link.get("packet_loss"), (int, float)) else str(link.get("packet_loss", "—")),
                    f"{link.get('latency_ms', 0):.1f} ms" if isinstance(link.get("latency_ms"), (int, float)) else str(link.get("latency_ms", "—")),
                    link.get("status", "—"),
                ])
            _add_data_table(doc, headers, rows, col_widths=[1.5, 1.5, 1.5, 1.0])

    # ── R-05 / R-08: Traffic Inbound ──────────────────────────────
    ti = rd.get("traffic_inbound", {})
    if ti:
        _add_heading_styled(doc, "Traffic Inbound", level=2)
        inbound_sites = ti.get("sites", {})
        if inbound_sites:
            for site_name, site_data in inbound_sites.items():
                _add_heading_styled(doc, f"Site: {site_name}", level=3)
                if site_data.get("top_services"):
                    _add_heading_styled(doc, "Top Inbound Services", level=4)
                    _add_hbar_items(doc, site_data["top_services"])

    # ── R-06 / R-08: Traffic Internal ─────────────────────────────
    tint = rd.get("traffic_internal", {})
    if tint:
        _add_heading_styled(doc, "Traffic Internal", level=2)
        int_sites = tint.get("sites", {})
        if int_sites:
            for site_name, site_data in int_sites.items():
                _add_heading_styled(doc, f"Site: {site_name}", level=3)
                if site_data.get("top_services"):
                    _add_heading_styled(doc, "Top Internal Services", level=4)
                    _add_hbar_items(doc, site_data["top_services"])
                if site_data.get("intra_lan_vs_inter_site"):
                    il = site_data["intra_lan_vs_inter_site"]
                    intra_val = f"{il.get('intra_lan_bytes', 0) / 1_000_000_000:.2f} GB" if il.get('intra_lan_bytes') else "0"
                    inter_val = f"{il.get('inter_site_bytes', 0) / 1_000_000_000:.2f} GB" if il.get('inter_site_bytes') else "0"
                    _add_kpi(doc, "Intra-LAN", intra_val)
                    _add_kpi(doc, "Inter-Site", inter_val)

    # ── R-07 / R-08: Executive Summary ────────────────────────────
    es = rd.get("executive_summary", {})
    if es:
        _add_heading_styled(doc, "Executive Summary", level=2)
        if es.get("narrative"):
            doc.add_paragraph(es["narrative"]).paragraph_format.space_after = Pt(6)

    # ── Footer (timestamp) ────────────────────────────────────────
    doc.add_paragraph("")  # spacer
    footer_p = doc.add_paragraph(
        f"Generated: {context.get('generated_at', '—')}  |  "
        f"{context.get('site_name', 'NOD')}  |  "
        f"{context.get('time_range', '—')}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = _hex_color(MUTED_HEX)
        run.font.name = "Calibri"

    # ── Save ──────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("DOCX report written: %s", output_path)
    return output_path

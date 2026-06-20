#!/usr/bin/env python3
"""
NOD Report DOCX Generator — generate_docx.py
==============================================
Programmatically builds a Word document (.docx) that visually matches
the NOD HTML/PDF reports as closely as python-docx allows.

Design System (from HTML/CSS):
  Typography : Calibri, Title=24pt, Section=18pt, Subtitle=14pt, Body=11pt, Small=9pt
  Colors     : Primary=#2563eb, Success=#10b981, Warning=#f59e0b, Danger=#ef4444,
               Text=#1e293b, Secondary=#64748b, Card-border=#e5e5e5
  Tables     : Header row blue (#dbeafe), 1pt borders, alternating white/#f8fafc
  Page       : A4, margins 20mm top/bottom, 15mm left/right
  Header     : NOD blue bar with white text (rendered via first-page header)
  Footer     : Page numbers, CONFIDENTIAL stamp (rendered via footer)

Usage:
  python generate_docx.py <context.json> <output.docx>

JSON Context Structure (same as the HTML template context):
{
  "report_meta": {
    "report_type": "R-01",
    "title": "Traffic Flow Report",
    "site": "Site_FGT-DC",
    "generated_at": "2025-07-14T10:00:00Z",
    "generated_by": "admin@nod.local",
    "job_id": "abc-123",
    "time_range": { "start": "...", "end": "..." }
  },
  "sections": {
    "kpi_summary": [...],
    "horizontal_bars": [...],
    "charts": [...],
    "tables": [...]
  },
  "charts": { "top_applications": "<base64>", ... },
  "devices": [...],
  ...
}
"""

from __future__ import annotations

import base64
import json
import logging
import sys
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

# ---------------------------------------------------------------------------
# python-docx imports
# ---------------------------------------------------------------------------
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import (
    Cm,
    Inches,
    Pt,
    RGBColor,
    Emu,
)

# ---------------------------------------------------------------------------
# Constants — NOD Design System
# ---------------------------------------------------------------------------
# ── Color Palette (matches HTML/CSS design tokens) ──
COLOR_PRIMARY       = RGBColor(0x25, 0x63, 0xEB)   # #2563eb
COLOR_PRIMARY_DARK  = RGBColor(0x1D, 0x4E, 0xD8)   # #1d4ed8
COLOR_PRIMARY_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)   # #dbeafe
COLOR_SUCCESS       = RGBColor(0x10, 0xB9, 0x81)   # #10b981
COLOR_WARNING       = RGBColor(0xF5, 0x9E, 0x0B)   # #f59e0b
COLOR_DANGER        = RGBColor(0xEF, 0x44, 0x44)   # #ef4444
COLOR_TEXT          = RGBColor(0x1E, 0x29, 0x3B)   # #1e293b
COLOR_SECONDARY     = RGBColor(0x64, 0x74, 0x8B)   # #64748b
COLOR_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)    # #ffffff
COLOR_CARD_BORDER   = RGBColor(0xE5, 0xE5, 0xE5)   # #e5e5e5
COLOR_ALT_ROW       = RGBColor(0xF8, 0xFA, 0xFC)   # #f8fafc
COLOR_LIGHT_GRAY    = RGBColor(0xF1, 0xF5, 0xF9)   # #f1f5f9
COLOR_BORDER_BLUE   = RGBColor(0xBF, 0xDB, 0xFE)   # lighter blue for table borders

# ── Typography Sizes ──
FONT_TITLE   = Pt(24)
FONT_SECTION = Pt(18)
FONT_SUB     = Pt(14)
FONT_BODY    = Pt(11)
FONT_SMALL   = Pt(9)
FONT_KPI_VAL = Pt(28)
FONT_KPI_LBL = Pt(9)

# ── Font Family ──
FONT_FAMILY = "Calibri"

# ── Layout ──
PAGE_TOP_MARGIN    = Cm(2.0)    # 20mm
PAGE_BOTTOM_MARGIN = Cm(2.0)
PAGE_LEFT_MARGIN   = Cm(1.5)    # 15mm
PAGE_RIGHT_MARGIN  = Cm(1.5)

logger = logging.getLogger(__name__)


# =========================================================================
# UTILITY HELPERS
# =========================================================================

def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert a hex color string (e.g. '#2563eb') to an RGBColor."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _set_cell_shading(cell, fill_color: RGBColor) -> None:
    """
    Set the background (shading) color of a table cell.
    Uses the underlying XML to apply a solid fill.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    # Remove any existing shading
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    # Create new shading element
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), f"{fill_color.red:02X}{fill_color.green:02X}{fill_color.blue:02X}")
    shading_elm.set(qn("w:val"), "clear")
    tc_pr.append(shading_elm)


def _set_cell_vertical_alignment(cell, alignment: str = "center") -> None:
    """Set vertical alignment of a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), alignment)
    tc_pr.append(v_align)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40, start: int = 80, end: int = 80) -> None:
    """
    Set cell margins (in twips: 1 inch = 1440 twips, 1mm ≈ 56.7 twips).
    Default: ~2mm top/bottom, ~4mm left/right.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    # Remove existing margins
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_cell_width(cell, width_cm: float) -> None:
    """Set explicit width on a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))  # cm to twips
    tc_w.set(qn("w:type"), "dxa")
    # Remove existing width
    for existing in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(existing)
    tc_pr.append(tc_w)


def _set_cell_no_border(cell) -> None:
    """Remove all borders from a cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tc_borders.append(border)
    # Remove existing borders
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    tc_pr.append(tc_borders)


def _set_table_borders(table, color: str = "E5E5E5", size: int = 4) -> None:
    """
    Set uniform 1pt borders on a table.
    size is in eighth-points: 4 = 0.5pt, 8 = 1pt
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Remove existing borders
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)

    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def _add_colored_paragraph(
    cell_or_paragraph,
    text: str,
    font_size: Pt = FONT_BODY,
    color: RGBColor = COLOR_TEXT,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_name: str = FONT_FAMILY,
    space_after: Pt = Pt(0),
    space_before: Pt = Pt(0),
):
    """
    Add a run with specific formatting to a paragraph.
    Returns the run for further manipulation.
    """
    # If we got a cell, get or create its paragraph
    if hasattr(cell_or_paragraph, "_tc"):
        para = cell_or_paragraph.paragraphs[0]
        if para.text:
            para = cell_or_paragraph.add_paragraph()
    else:
        para = cell_or_paragraph

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = color
    run.bold = bold
    para.alignment = alignment
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return run


def _add_page_break(doc: Document) -> None:
    """Insert a hard page break."""
    doc.add_page_break()


def _set_table_layout_autofit(table) -> None:
    """Make the table use fixed column widths (autofit off)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    # Remove existing layout
    for existing in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_layout)


def _create_horizontal_bar_in_cell(
    cell,
    label: str,
    value: str,
    percentage: float,          # 0.0 – 1.0
    bar_color: RGBColor,
    rank: Optional[int] = None,
) -> None:
    """
    Render a single horizontal bar chart row inside a single table cell.
    Uses a nested table: [rank | label | ████░░░░ | value]

    This is the Word equivalent of the HTML horizontal bar chart component.
    """
    # Clear existing content
    for p in cell.paragraphs:
        p.clear()

    # We'll build this as a formatted paragraph line with Unicode block chars
    # for the progress bar, since nested tables inside cells can be complex.
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Rank number
    if rank is not None:
        rank_run = para.add_run(f" {rank}  ")
        rank_run.font.size = Pt(10)
        rank_run.font.color.rgb = COLOR_SECONDARY
        rank_run.bold = True
        rank_run.font.name = FONT_FAMILY

    # Label
    label_run = para.add_run(f"{label:<25}")
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = COLOR_TEXT
    label_run.font.name = FONT_FAMILY

    # Bar representation using Unicode block characters
    bar_width = 20  # number of block characters
    filled = max(1, int(percentage * bar_width))
    empty = bar_width - filled
    bar_text = "█" * filled + "░" * empty
    bar_run = para.add_run(bar_text)
    bar_run.font.size = Pt(8)
    bar_run.font.color.rgb = bar_color
    bar_run.font.name = "Consolas"  # monospace for alignment

    # Value
    value_run = para.add_run(f"  {value}")
    value_run.font.size = Pt(10)
    value_run.font.color.rgb = COLOR_TEXT
    value_run.bold = True
    value_run.font.name = FONT_FAMILY


# =========================================================================
# PAGE SETUP
# =========================================================================

def _setup_page(doc: Document) -> None:
    """
    Configure A4 page dimensions and margins to match the HTML/CSS design:
      - A4 portrait
      - Top/Bottom: 20mm, Left/Right: 15mm
    """
    section = doc.sections[0]
    section.page_width  = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin    = PAGE_TOP_MARGIN
    section.bottom_margin = PAGE_BOTTOM_MARGIN
    section.left_margin   = PAGE_LEFT_MARGIN
    section.right_margin  = PAGE_RIGHT_MARGIN


# =========================================================================
# HEADER / FOOTER
# =========================================================================

def _setup_header_footer(doc: Document, context: dict) -> None:
    """
    Add a branded header and footer to all sections.

    Header: Blue bar with NOD branding (matches HTML .nod-header)
    Footer: Page number + CONFIDENTIAL stamp (matches HTML .nod-footer)
    """
    section = doc.sections[0]
    report_meta = context.get("report_meta", {})
    report_type = report_meta.get("report_type", "NOD")
    title = report_meta.get("title", "NOD Report")

    # ── HEADER ──
    header = section.header
    header.is_linked_to_previous = False

    # Main header table (blue background bar)
    hdr_table = header.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style the header table
    _set_table_borders(hdr_table, color="2563EB", size=4)

    # Left cell — NOD brand + title
    left_cell = hdr_table.cell(0, 0)
    _set_cell_shading(left_cell, COLOR_PRIMARY)
    _set_cell_no_border(left_cell)
    _set_cell_margins(left_cell, top=60, bottom=60, start=100, end=100)

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    nod_run = left_para.add_run("NOD")
    nod_run.font.size = Pt(14)
    nod_run.font.bold = True
    nod_run.font.color.rgb = COLOR_WHITE
    nod_run.font.name = FONT_FAMILY

    sep_run = left_para.add_run("  │  ")
    sep_run.font.size = Pt(11)
    sep_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sep_run.font.name = FONT_FAMILY

    title_run = left_para.add_run(f"{title}")
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = COLOR_WHITE
    title_run.font.name = FONT_FAMILY

    # Right cell — report type badge
    right_cell = hdr_table.cell(0, 1)
    _set_cell_shading(right_cell, COLOR_PRIMARY)
    _set_cell_no_border(right_cell)
    _set_cell_margins(right_cell, top=60, bottom=60, start=100, end=100)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    badge_run = right_para.add_run(report_type)
    badge_run.font.size = Pt(12)
    badge_run.font.bold = True
    badge_run.font.color.rgb = COLOR_WHITE
    badge_run.font.name = FONT_FAMILY

    # Set column widths: 70% left, 30% right
    left_cell.width  = Cm(12.0)
    right_cell.width = Cm(6.0)

    # ── FOOTER ──
    footer = section.footer
    footer.is_linked_to_previous = False

    # Separator line (thin table)
    ftr_table = footer.add_table(rows=1, cols=1, width=section.page_width - section.left_margin - section.right_margin)
    ftr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    line_cell = ftr_table.cell(0, 0)
    _set_cell_margins(line_cell, top=0, bottom=0, start=0, end=0)

    line_para = line_cell.paragraphs[0]
    # Add a horizontal line using a paragraph border
    pPr = line_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E2E8F0")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer text: brand + page number + confidential
    ftr_para = footer.add_paragraph()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr_para.paragraph_format.space_before = Pt(4)

    brand_run = ftr_para.add_run("NOD — Network Observability Dashboard")
    brand_run.font.size = Pt(8)
    brand_run.font.color.rgb = COLOR_PRIMARY
    brand_run.font.bold = True
    brand_run.font.name = FONT_FAMILY

    sep1 = ftr_para.add_run("  |  ")
    sep1.font.size = Pt(8)
    sep1.font.color.rgb = COLOR_SECONDARY

    type_run = ftr_para.add_run(f"Report: {report_type}")
    type_run.font.size = Pt(8)
    type_run.font.color.rgb = COLOR_SECONDARY
    type_run.font.name = FONT_FAMILY

    # Page number field
    sep2 = ftr_para.add_run("  |  Page ")
    sep2.font.size = Pt(8)
    sep2.font.color.rgb = COLOR_SECONDARY

    # Insert PAGE field
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    page_run = ftr_para.add_run()
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = COLOR_SECONDARY
    page_run._r.append(fldChar_begin)
    page_run._r.append(instrText)
    page_run._r.append(fldChar_end)

    # CONFIDENTIAL line
    conf_para = footer.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_before = Pt(2)
    conf_run = conf_para.add_run("DOCUMENT CLASSIFICATION: INTERNAL — CONFIDENTIAL")
    conf_run.font.size = Pt(7)
    conf_run.font.color.rgb = COLOR_DANGER
    conf_run.font.bold = True
    conf_run.font.name = FONT_FAMILY


# =========================================================================
# TITLE PAGE
# =========================================================================

def _build_title_page(doc: Document, context: dict) -> None:
    """
    Create a professional title page matching the HTML header design:
      - Blue background banner
      - NOD logo/icon
      - Report title
      - Subtitle
      - Metadata bar (generated, author, job ID, time range)
    """
    report_meta = context.get("report_meta", {})
    title = report_meta.get("title", "NOD Report")
    report_type = report_meta.get("report_type", "R-XX")
    generated_at = report_meta.get("generated_at", "")
    generated_by = report_meta.get("generated_by", "")
    job_id = report_meta.get("job_id", "")
    time_range = report_meta.get("time_range", {})
    time_range_str = ""
    if isinstance(time_range, dict):
        time_range_str = f"{time_range.get('start', 'N/A')} — {time_range.get('end', 'N/A')}"
    elif isinstance(time_range, str):
        time_range_str = time_range

    # ── Blue Banner Table (simulates the HTML .nod-header) ──
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table to full page width
    tbl = banner_table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)

    # Remove default borders from banner
    _set_table_borders(banner_table, color="2563EB", size=4)

    banner_cell = banner_table.cell(0, 0)
    _set_cell_shading(banner_cell, COLOR_PRIMARY)
    _set_cell_no_border(banner_cell)
    _set_cell_margins(banner_cell, top=200, bottom=200, start=200, end=200)

    # NOD Logo text
    logo_para = banner_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after = Pt(4)

    logo_icon = logo_para.add_run("  NOD  ")
    logo_icon.font.size = Pt(22)
    logo_icon.font.bold = True
    logo_icon.font.color.rgb = COLOR_WHITE
    logo_icon.font.name = FONT_FAMILY

    logo_sep = logo_para.add_run("  │  ")
    logo_sep.font.size = Pt(16)
    logo_sep.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)  # blue-300

    brand_text = logo_para.add_run("Network Observability Dashboard")
    brand_text.font.size = Pt(11)
    brand_text.font.bold = True
    brand_text.font.color.rgb = COLOR_WHITE
    brand_text.font.name = FONT_FAMILY
    brand_text.font.all_caps = True

    # Spacer
    spacer = banner_cell.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    spacer.paragraph_format.space_after = Pt(0)

    # Report Title (large)
    h1_para = banner_cell.add_paragraph()
    h1_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_para.paragraph_format.space_after = Pt(4)
    h1_run = h1_para.add_run(title)
    h1_run.font.size = FONT_TITLE
    h1_run.font.bold = True
    h1_run.font.color.rgb = COLOR_WHITE
    h1_run.font.name = FONT_FAMILY

    # Subtitle
    sub_para = banner_cell.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_para.paragraph_format.space_after = Pt(0)
    sub_run = sub_para.add_run("Comprehensive Network Intelligence Report")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0xDB, 0xEA, 0xFE)  # blue-100
    sub_run.font.name = FONT_FAMILY

    # ── Metadata Bar (matches HTML .meta-bar) ──
    doc.add_paragraph()  # spacer
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style meta bar
    _set_table_borders(meta_table, color="E2E8F0", size=4)

    meta_items = [
        ("Generated", str(generated_at)[:19] if generated_at else "N/A"),
        ("Author", str(generated_by)),
        ("Job ID", str(job_id)[:12] if job_id else "N/A"),
        ("Time Range", time_range_str[:40] if time_range_str else "N/A"),
    ]

    for idx, (label, value) in enumerate(meta_items):
        cell = meta_table.cell(0, idx)
        _set_cell_shading(cell, COLOR_WHITE)
        _set_cell_margins(cell, top=40, bottom=40, start=60, end=60)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        lbl_run = para.add_run(f"{label}: ")
        lbl_run.font.size = Pt(8)
        lbl_run.font.bold = True
        lbl_run.font.color.rgb = COLOR_TEXT
        lbl_run.font.name = FONT_FAMILY

        val_run = para.add_run(value)
        val_run.font.size = Pt(8)
        val_run.font.color.rgb = COLOR_SECONDARY
        val_run.font.name = FONT_FAMILY

    # Page break after title page
    _add_page_break(doc)


# =========================================================================
# SECTION HEADINGS
# =========================================================================

def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    """
    Add a styled section heading that matches the HTML design:
      - Level 1: 18pt, bold, blue bottom border
      - Level 2: 14pt, bold
      - Level 3: 11pt, bold
    """
    heading = doc.add_heading(text, level=level)

    # Style the heading run
    for run in heading.runs:
        run.font.name = FONT_FAMILY
        run.font.color.rgb = COLOR_TEXT

        if level == 1:
            run.font.size = FONT_SECTION
            run.font.bold = True
        elif level == 2:
            run.font.size = FONT_SUB
            run.font.bold = True
        else:
            run.font.size = FONT_BODY
            run.font.bold = True

    # Add blue bottom border for level 1 headings (matches CSS border-bottom: 3px solid primary)
    if level == 1:
        pPr = heading._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")  # ~1.5pt
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2563EB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    heading.paragraph_format.space_before = Pt(24 if level == 1 else 12)
    heading.paragraph_format.space_after = Pt(8)


# =========================================================================
# KPI CARDS
# =========================================================================

def _build_kpi_section(doc: Document, kpi_items: list[dict]) -> None:
    """
    Build KPI summary cards in a table grid layout.

    Each KPI item:
      { "label": "Active VPN", "value": "1,240", "status": "success|warning|danger|info" }

    Layout: Table with 3-4 columns, each cell = one KPI card.
    Styling matches HTML .kpi-card: centered text, colored value, uppercase label.
    """
    if not kpi_items:
        return

    _add_section_heading(doc, "Key Performance Indicators", level=1)

    # Determine grid: 3 columns for ≤6 items, 4 for >6
    cols = 4 if len(kpi_items) > 6 else 3
    rows_needed = (len(kpi_items) + cols - 1) // cols

    kpi_table = doc.add_table(rows=rows_needed, cols=cols)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style the table: no visible borders, centered cells
    _set_table_borders(kpi_table, color="E5E5E5", size=4)

    for idx, item in enumerate(kpi_items):
        row_idx = idx // cols
        col_idx = idx % cols
        cell = kpi_table.cell(row_idx, col_idx)

        # White background card
        _set_cell_shading(cell, COLOR_WHITE)
        _set_cell_margins(cell, top=100, bottom=100, start=80, end=80)
        _set_cell_vertical_alignment(cell, "center")

        # Clear default paragraph
        cell.paragraphs[0].clear()

        # Label (uppercase, small, gray)
        label_para = cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.paragraph_format.space_after = Pt(4)
        label_para.paragraph_format.space_before = Pt(0)
        lbl_run = label_para.add_run(item.get("label", "").upper())
        lbl_run.font.size = FONT_KPI_LBL
        lbl_run.font.bold = True
        lbl_run.font.color.rgb = COLOR_SECONDARY
        lbl_run.font.name = FONT_FAMILY

        # Value (large, colored by status)
        value_para = cell.add_paragraph()
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_para.paragraph_format.space_before = Pt(2)
        value_para.paragraph_format.space_after = Pt(0)

        status = item.get("status", "info")
        value_color = {
            "success": COLOR_SUCCESS,
            "warning": COLOR_WARNING,
            "danger": COLOR_DANGER,
            "info": COLOR_PRIMARY,
        }.get(status, COLOR_PRIMARY)

        val_run = value_para.add_run(str(item.get("value", "—")))
        val_run.font.size = FONT_KPI_VAL
        val_run.font.bold = True
        val_run.font.color.rgb = value_color
        val_run.font.name = FONT_FAMILY

    # Fill empty cells with white background and no borders
    total_cells = rows_needed * cols
    for idx in range(len(kpi_items), total_cells):
        row_idx = idx // cols
        col_idx = idx % cols
        cell = kpi_table.cell(row_idx, col_idx)
        _set_cell_shading(cell, COLOR_WHITE)
        _set_cell_no_border(cell)


# =========================================================================
# HORIZONTAL BAR CHARTS
# =========================================================================

def _build_horizontal_bars(doc: Document, bar_sections: list[dict]) -> None:
    """
    Build horizontal bar chart sections.

    Each bar section:
      {
        "title": "Top Applications",
        "items": [
          { "rank": 1, "name": "HTTPS", "value": "45%", "color": "#2563eb" }
        ]
      }

    Layout: One table per section with columns for each bar item.
    Uses Unicode block characters for the bar visualization.
    """
    if not bar_sections:
        return

    _add_section_heading(doc, "Rankings & Comparisons", level=1)

    for section in bar_sections:
        title = section.get("title", "")
        items = section.get("items", [])

        if not items:
            continue

        # Section subtitle
        _add_section_heading(doc, title, level=2)

        if not items:
            continue

        # Parse max value for percentage calculation
        max_val = 0
        parsed_items = []
        for item in items:
            raw_val = item.get("value", "0")
            # Try to parse numeric value (strip %, MB, GB etc.)
            numeric_str = ""
            for ch in str(raw_val):
                if ch.isdigit() or ch == ".":
                    numeric_str += ch
                else:
                    break
            try:
                numeric_val = float(numeric_str) if numeric_str else 0
            except ValueError:
                numeric_val = 0
            parsed_items.append({**item, "_numeric": numeric_val})
            max_val = max(max_val, numeric_val)

        # Create a table: 1 column per bar item (or use rows)
        # Better approach: use a single-column table with one row per bar
        bar_table = doc.add_table(rows=len(parsed_items), cols=1)
        bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        _set_table_borders(bar_table, color="E5E5E5", size=4)

        for row_idx, item in enumerate(parsed_items):
            cell = bar_table.cell(row_idx, 0)

            # Alternating row background
            if row_idx % 2 == 1:
                _set_cell_shading(cell, COLOR_ALT_ROW)
            else:
                _set_cell_shading(cell, COLOR_WHITE)

            _set_cell_margins(cell, top=40, bottom=40, start=80, end=80)

            # Parse color
            item_color_hex = item.get("color", "#2563eb")
            item_color = _hex_to_rgb(item_color_hex)

            # Calculate percentage for bar width
            pct = (item["_numeric"] / max_val) if max_val > 0 else 0

            _create_horizontal_bar_in_cell(
                cell=cell,
                label=item.get("name", ""),
                value=str(item.get("value", "")),
                percentage=pct,
                bar_color=item_color,
                rank=item.get("rank"),
            )

        # Spacer after bar chart
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)


# =========================================================================
# CHART IMAGES
# =========================================================================

def _build_chart_section(doc: Document, charts: list[dict]) -> None:
    """
    Embed chart images from base64-encoded PNGs.

    Each chart:
      { "title": "Throughput Over Time", "image_base64": "<base64-encoded PNG>" }
    """
    if not charts:
        return

    _add_section_heading(doc, "Charts & Visualizations", level=1)

    for chart in charts:
        title = chart.get("title", "Chart")
        img_b64 = chart.get("image_base64", "")

        if not img_b64:
            continue

        # Chart title
        _add_section_heading(doc, title, level=2)

        # Decode and embed image
        try:
            img_bytes = base64.b64decode(img_b64)
            img_stream = BytesIO(img_bytes)

            # Add image centered, scaled to fit page width
            # Page usable width ≈ 21cm - 3cm margins = 18cm = ~7.09 inches
            max_width = Inches(6.8)  # slightly less than full width for padding
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(8)

            run = para.add_run()
            run.add_picture(img_stream, width=max_width)

        except Exception as e:
            logger.warning("Failed to embed chart image '%s': %s", title, e)
            err_para = doc.add_paragraph()
            err_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            err_run = err_para.add_run(f"[Chart image unavailable: {title}]")
            err_run.font.color.rgb = COLOR_SECONDARY
            err_run.font.size = FONT_SMALL
            err_run.font.italic = True


# =========================================================================
# BASE64 CHART IMAGES (from context["charts"] dict)
# =========================================================================

def _build_embedded_charts(doc: Document, charts_dict: dict) -> None:
    """
    Embed chart images from the flat dict format used by the backend:
      context["charts"] = { "top_applications": "<base64>", ... }

    This handles the charts produced by chart_renderer.py.
    """
    if not charts_dict:
        return

    _add_section_heading(doc, "Charts & Visualizations", level=1)

    # Map of chart keys to display titles
    chart_titles = {
        "top_applications": "Top 10 Applications by Traffic Volume",
        "throughput": "Throughput Over Time",
        "top_as_orgs": "Top 10 Destination AS Organizations",
        "top_countries": "Top 10 Destination Countries",
        "protocol_dist": "Protocol Distribution by Traffic Volume",
        "cpu_timeline": "CPU Usage Over Time",
        "vpn_users": "VPN Users Overview",
    }

    for chart_key, img_b64 in charts_dict.items():
        title = chart_titles.get(str(chart_key), str(chart_key).replace("_", " ").title())

        if not img_b64:
            continue

        # Chart title
        _add_section_heading(doc, title, level=2)

        try:
            img_bytes = base64.b64decode(img_b64)
            img_stream = BytesIO(img_bytes)

            # Add image centered, scaled to fit page width
            max_width = Inches(6.8)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(8)

            run = para.add_run()
            run.add_picture(img_stream, width=max_width)

        except Exception as e:
            logger.warning("Failed to embed chart '%s': %s", chart_key, e)
            err_para = doc.add_paragraph()
            err_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            err_run = err_para.add_run(f"[Chart image unavailable: {chart_key}]")
            err_run.font.color.rgb = COLOR_SECONDARY
            err_run.font.size = FONT_SMALL
            err_run.font.italic = True


# =========================================================================
# DATA TABLES
# =========================================================================

def _build_data_table(
    doc: Document,
    title: str,
    columns: list[str],
    rows: list[list[str]],
    highlight_columns: Optional[dict[int, str]] = None,
) -> None:
    """
    Build a styled data table matching the HTML table design:
      - Header row: blue background (#dbeafe), bold dark-blue text
      - Alternating row colors: white / #f8fafc
      - 1pt borders (#e5e5e5)
      - Optional column-specific styling (e.g., status colors)

    highlight_columns: dict mapping column_index -> "green"|"amber"|"red"|"blue"
    """
    if not columns or not rows:
        return

    # Section title
    _add_section_heading(doc, title, level=2)

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style table borders
    _set_table_borders(table, color="E5E5E5", size=4)

    # Header row
    for col_idx, col_name in enumerate(columns):
        cell = table.cell(0, col_idx)
        _set_cell_shading(cell, COLOR_PRIMARY_LIGHT)
        _set_cell_margins(cell, top=60, bottom=60, start=80, end=80)
        _set_cell_vertical_alignment(cell, "center")

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)

        run = para.add_run(col_name.upper())
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY_DARK
        run.font.name = FONT_FAMILY

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_value in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)

            # Alternating row background
            if row_idx % 2 == 1:
                _set_cell_shading(cell, COLOR_ALT_ROW)
            else:
                _set_cell_shading(cell, COLOR_WHITE)

            _set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            _set_cell_vertical_alignment(cell, "center")

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)

            run = para.add_run(str(cell_value))
            run.font.size = Pt(9)
            run.font.name = FONT_FAMILY

            # Apply status-based coloring if this column has it
            if highlight_columns and col_idx in highlight_columns:
                status = highlight_columns[col_idx]
                color = {
                    "green": COLOR_SUCCESS,
                    "amber": COLOR_WARNING,
                    "red": COLOR_DANGER,
                    "blue": COLOR_PRIMARY,
                }.get(status, COLOR_TEXT)
                run.font.color.rgb = color
                run.font.bold = True
            else:
                run.font.color.rgb = COLOR_TEXT


# =========================================================================
# REPORT SECTION BUILDERS (one per report type)
# =========================================================================

def _build_r01_traffic_flow(doc: Document, context: dict) -> None:
    """R-01: Traffic Flow Report — charts, per-site tables, protocol dist."""
    charts = context.get("charts", {})
    has_traffic_charts = any(
        charts.get(k) for k in
        ("top_applications", "throughput", "top_as_orgs", "top_countries", "protocol_dist")
    )
    per_site = context.get("per_site_traffic", {})
    total_tp = context.get("total_throughput_bytes", 0)

    if not has_traffic_charts and not per_site and not total_tp:
        return

    _add_section_heading(doc, "R-01: Traffic Flow Analysis", level=1)

    # Total throughput KPI
    if total_tp:
        kpi_items = [
            {"label": "Total Throughput", "value": f"{total_tp:,} bytes", "status": "info"},
        ]
        _build_kpi_section(doc, kpi_items)

    # Embedded charts
    traffic_charts = {k: v for k, v in charts.items() if k in
                      ("top_applications", "throughput", "top_as_orgs", "top_countries", "protocol_dist")}
    if traffic_charts:
        _build_embedded_charts(doc, traffic_charts)

    # Per-site traffic summary table
    if per_site:
        _add_section_heading(doc, "Per-Site Traffic Summary", level=2)
        columns = ["Site", "Total Bytes", "Top Application", "Sessions"]
        rows = []
        for site, site_data in per_site.items():
            top_app = "N/A"
            if isinstance(site_data, dict):
                top_apps = site_data.get("top_apps", [])
                if top_apps and isinstance(top_apps, list):
                    top_app = top_apps[0].get("app_name", "N/A")
            rows.append([
                site,
                str(site_data.get("total_bytes", "N/A")) if isinstance(site_data, dict) else "N/A",
                top_app,
                str(site_data.get("total_sessions", "N/A")) if isinstance(site_data, dict) else "N/A",
            ])
        _build_data_table(doc, "Per-Site Traffic Summary", columns, rows)


def _build_r02_resource_usage(doc: Document, context: dict) -> None:
    """R-02: Resource Usage Report — device table, CPU chart."""
    devices = context.get("devices", [])
    charts = context.get("charts", {})

    if not devices:
        return

    _add_section_heading(doc, "R-02: Device Resource Status", level=1)

    # Device table
    if devices:
        columns = ["Device", "Hostname", "CPU %", "Memory %", "Sessions", "Sync Status"]
        rows = []
        for d in devices:
            cpu = str(d.get("cpu_usage", ""))
            mem = str(d.get("mem_usage", ""))
            rows.append([
                d.get("device", ""),
                d.get("hostname", ""),
                cpu,
                mem,
                str(d.get("session_count", "")),
                d.get("sync_status", ""),
            ])

        # Build table with status-based coloring for CPU and Memory columns
        _build_data_table(
            doc,
            "Device Resource Status",
            columns,
            rows,
            highlight_columns={},  # We'll handle inline below
        )

        # Re-style the table with status colors by scanning the actual values
        # (This is a second pass to apply conditional coloring)
        # Find the most recently added table
        if doc.tables:
            last_table = doc.tables[-1]
            for row_idx in range(1, len(last_table.rows)):
                d = devices[row_idx - 1] if row_idx - 1 < len(devices) else {}
                cpu_val = d.get("cpu_usage", 0)
                mem_val = d.get("mem_usage", 0)

                # Color CPU cell (col 2)
                cpu_cell = last_table.cell(row_idx, 2)
                cpu_color = (
                    COLOR_SUCCESS if cpu_val < 60
                    else COLOR_WARNING if cpu_val < 80
                    else COLOR_DANGER
                )
                for para in cpu_cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = cpu_color
                        run.font.bold = True

                # Color Memory cell (col 3)
                mem_cell = last_table.cell(row_idx, 3)
                mem_color = (
                    COLOR_SUCCESS if mem_val < 60
                    else COLOR_WARNING if mem_val < 80
                    else COLOR_DANGER
                )
                for para in mem_cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = mem_color
                        run.font.bold = True

    # CPU timeline chart
    cpu_chart = charts.get("cpu_timeline")
    if cpu_chart:
        _build_embedded_charts(doc, {"cpu_timeline": cpu_chart})


def _build_r03_vpn_users(doc: Document, context: dict) -> None:
    """R-03: Active VPN Users Report — VPN counts, chart."""
    ssl_vpn = context.get("ssl_vpn_count", 0)
    ipsec_vpn = context.get("ipsec_vpn_count", 0)
    charts = context.get("charts", {})

    if not ssl_vpn and not ipsec_vpn and not charts.get("vpn_users"):
        return

    _add_section_heading(doc, "R-03: Active VPN Users", level=1)

    # KPI cards
    kpi_items = [
        {"label": "SSL VPN Users", "value": str(ssl_vpn or 0), "status": "success"},
        {"label": "IPsec VPN Users", "value": str(ipsec_vpn or 0), "status": "info"},
    ]
    _build_kpi_section(doc, kpi_items)

    # VPN bar chart
    if charts.get("vpn_users"):
        _build_embedded_charts(doc, {"vpn_users": charts["vpn_users"]})


def _build_r04_sdwan_sla(doc: Document, context: dict) -> None:
    """R-04: SD-WAN SLA Report — SLA charts, compliance table."""
    sla_data = context.get("sla_data", {})
    sla_charts = context.get("sla_charts", {})

    if not sla_data and not sla_charts:
        return

    _add_section_heading(doc, "R-04: SD-WAN SLA Monitoring", level=1)

    # SLA timeline charts
    if sla_charts:
        _build_embedded_charts(doc, sla_charts)

    # SLA compliance summary table
    if sla_data:
        _add_section_heading(doc, "SLA Compliance Summary", level=2)
        columns = ["Site", "Avg Latency (ms)", "Avg Jitter (ms)", "Avg Packet Loss (%)"]
        rows = []
        for site, site_sla in sla_data.items():
            summary = site_sla.get("summary", {})
            if summary:
                rows.append([
                    site,
                    str(summary.get("avg_latency", "N/A")),
                    str(summary.get("avg_jitter", "N/A")),
                    str(summary.get("avg_packet_loss", "N/A")),
                ])
        if rows:
            _build_data_table(doc, "SLA Compliance Summary", columns, rows)


def _build_r05_traffic_inbound(doc: Document, context: dict) -> None:
    """R-05: Traffic Inbound Report — inbound charts, per-site tables."""
    inbound = context.get("inbound_data", {})
    charts = context.get("charts", {})

    if not inbound and not charts:
        return

    _add_section_heading(doc, "R-05: Traffic Inbound Analysis", level=1)

    # Inbound charts
    inbound_charts = {k: v for k, v in charts.items() if k.startswith("inbound_")}
    if inbound_charts:
        _build_embedded_charts(doc, inbound_charts)

    # Per-site inbound service tables
    for site, site_data in (inbound.items() if isinstance(inbound, dict) else []):
        if not isinstance(site_data, dict):
            continue
        top_svcs = site_data.get("top_services", [])
        if top_svcs:
            _add_section_heading(doc, f"{site} — Top Inbound Services", level=2)
            columns = ["#", "Service", "Total Bytes"]
            rows = [
                [str(i + 1), s.get("service_name", ""), str(s.get("total_bytes", ""))]
                for i, s in enumerate(top_svcs[:10])
            ]
            _build_data_table(doc, f"{site} Services", columns, rows)


def _build_r06_traffic_internal(doc: Document, context: dict) -> None:
    """R-06: Traffic Internal Report — internal per-site tables."""
    internal = context.get("internal_data", {})

    if not internal:
        return

    _add_section_heading(doc, "R-06: Traffic Internal Analysis", level=1)

    for site, site_data in (internal.items() if isinstance(internal, dict) else []):
        if not isinstance(site_data, dict):
            continue
        top_svcs = site_data.get("top_services", [])
        if top_svcs:
            _add_section_heading(doc, f"{site} — Top Internal Services", level=2)
            columns = ["#", "Service", "Total Bytes"]
            rows = [
                [str(i + 1), s.get("service_name", ""), str(s.get("total_bytes", ""))]
                for i, s in enumerate(top_svcs[:10])
            ]
            _build_data_table(doc, f"{site} Internal Services", columns, rows)


def _build_r07_executive_summary(doc: Document, context: dict) -> None:
    """R-07: Executive Summary — top-level KPIs, top apps, device health."""
    exec_summary = context.get("executive_summary", {})
    total_tp = context.get("total_throughput_bytes", 0)
    ssl_vpn = context.get("ssl_vpn_count", 0)
    ipsec_vpn = context.get("ipsec_vpn_count", 0)
    devices = context.get("devices", [])
    sla_data = context.get("sla_data", {})

    # Build executive KPIs
    kpi_items = []
    if total_tp:
        kpi_items.append({"label": "Total Throughput", "value": f"{total_tp:,}", "status": "info"})
    kpi_items.append({"label": "SSL VPN Users", "value": str(ssl_vpn or 0), "status": "success"})
    kpi_items.append({"label": "IPsec VPN Users", "value": str(ipsec_vpn or 0), "status": "info"})
    if devices:
        healthy = sum(1 for d in devices if d.get("cpu_usage", 0) < 80 and d.get("mem_usage", 0) < 80)
        kpi_items.append({"label": "Total Devices", "value": str(len(devices)), "status": "info"})
        kpi_items.append({"label": "Healthy Devices", "value": str(healthy), "status": "success"})
        kpi_items.append({"label": "Degraded Devices", "value": str(len(devices) - healthy), "status": "warning"})
    if sla_data:
        kpi_items.append({"label": "SLA Sites", "value": str(len(sla_data)), "status": "info"})

    if kpi_items:
        _add_section_heading(doc, "R-07: Executive Summary", level=1)
        _build_kpi_section(doc, kpi_items)

    # Top 5 apps table
    top_5 = exec_summary.get("top_5_apps", []) if isinstance(exec_summary, dict) else []
    if not top_5:
        top_apps = context.get("top_applications", [])
        if top_apps:
            top_5 = [{"app": a.get("application", ""), "bytes": a.get("total_bytes", 0)} for a in top_apps[:5]]

    if top_5:
        _add_section_heading(doc, "Top 5 Applications", level=2)
        columns = ["#", "Application", "Total Bytes"]
        rows = [[str(i + 1), a.get("app", ""), str(a.get("bytes", ""))] for i, a in enumerate(top_5)]
        _build_data_table(doc, "Top 5 Applications", columns, rows)

    # Device health overview
    if devices:
        _add_section_heading(doc, "Device Health Overview", level=2)
        columns = ["Device", "CPU %", "Memory %", "Status"]
        rows = []
        for d in devices:
            cpu = d.get("cpu_usage", 0)
            mem = d.get("mem_usage", 0)
            status = "Healthy" if cpu < 80 and mem < 80 else "Degraded"
            rows.append([d.get("device", ""), str(cpu), str(mem), status])
        _build_data_table(doc, "Device Health", columns, rows)

    # SLA summary
    if sla_data:
        _add_section_heading(doc, "SLA Summary", level=2)
        columns = ["Site", "Avg Latency (ms)", "Avg Jitter (ms)", "Avg Packet Loss (%)"]
        rows = []
        for site, site_sla in sla_data.items():
            summary = site_sla.get("summary", {})
            if summary:
                rows.append([
                    site,
                    str(summary.get("avg_latency", "N/A")),
                    str(summary.get("avg_jitter", "N/A")),
                    str(summary.get("avg_packet_loss", "N/A")),
                ])
        if rows:
            _build_data_table(doc, "SLA Summary", columns, rows)


# =========================================================================
# GENERIC SECTION BUILDER (from JSON context sections)
# =========================================================================

def _build_generic_sections(doc: Document, sections: dict) -> None:
    """
    Build sections from the generic JSON structure:
      "sections": {
        "kpi_summary": [...],
        "horizontal_bars": [...],
        "charts": [...],
        "tables": [...]
      }
    """
    # KPI Summary
    kpi_items = sections.get("kpi_summary", [])
    if kpi_items:
        _build_kpi_section(doc, kpi_items)

    # Horizontal Bars
    bar_sections = sections.get("horizontal_bars", [])
    if bar_sections:
        _build_horizontal_bars(doc, bar_sections)

    # Charts (from sections.charts — list format)
    section_charts = sections.get("charts", [])
    if section_charts:
        _build_chart_section(doc, section_charts)

    # Tables
    section_tables = sections.get("tables", [])
    for tbl in section_tables:
        title = tbl.get("title", "Data")
        columns = tbl.get("columns", [])
        rows = tbl.get("rows", [])
        if columns and rows:
            _build_data_table(doc, title, columns, rows)


# =========================================================================
# DOCUMENT CLASSIFICATION PAGE
# =========================================================================

def _build_classification_page(doc: Document, context: dict) -> None:
    """
    Add a document classification page (matches HTML footer).
    Used as the final page of R-08 (All-in-One) reports.
    """
    _add_page_break(doc)

    # Classification heading
    heading = doc.add_heading("Document Classification", level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_TEXT
        run.font.name = FONT_FAMILY

    report_meta = context.get("report_meta", {})
    report_title = report_meta.get("title", "NOD Report")
    job_id = report_meta.get("job_id", "")
    generated_at = report_meta.get("generated_at", "")
    generated_by = report_meta.get("generated_by", "")

    # Classification table
    cls_table = doc.add_table(rows=5, cols=2)
    cls_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(cls_table, color="E5E5E5", size=4)

    cls_data = [
        ("Classification", "INTERNAL — CONFIDENTIAL"),
        ("Report Type", report_title),
        ("Job ID", str(job_id)),
        ("Generated", str(generated_at)),
        ("Author", str(generated_by)),
    ]

    for idx, (label, value) in enumerate(cls_data):
        # Label cell
        label_cell = cls_table.cell(idx, 0)
        _set_cell_shading(label_cell, COLOR_PRIMARY_LIGHT)
        _set_cell_margins(label_cell, top=60, bottom=60, start=80, end=80)

        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.size = Pt(10)
        label_run.font.bold = True
        label_run.font.color.rgb = COLOR_PRIMARY_DARK
        label_run.font.name = FONT_FAMILY

        # Value cell
        value_cell = cls_table.cell(idx, 1)
        _set_cell_shading(value_cell, COLOR_WHITE)
        _set_cell_margins(value_cell, top=60, bottom=60, start=80, end=80)

        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        value_run.font.size = Pt(10)
        value_run.font.color.rgb = COLOR_TEXT
        value_run.font.name = FONT_FAMILY

        # Highlight confidential classification
        if label == "Classification":
            value_run.font.color.rgb = COLOR_DANGER
            value_run.font.bold = True

    # Set column widths
    for row in cls_table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(13.0)

    # Footer stamp
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(24)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(12)

    stamp = footer_para.add_run("NOD — Network Observability Dashboard")
    stamp.font.size = Pt(11)
    stamp.font.bold = True
    stamp.font.color.rgb = COLOR_PRIMARY
    stamp.font.name = FONT_FAMILY

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_before = Pt(8)
    conf_run = conf_para.add_run("Document Classification: Internal — Confidential")
    conf_run.font.size = Pt(10)
    conf_run.font.bold = True
    conf_run.font.color.rgb = COLOR_DANGER
    conf_run.font.name = FONT_FAMILY


# =========================================================================
# MAIN ENTRY POINT
# =========================================================================

def generate_docx_report(context: dict, output_path: str | Path) -> Path:
    """
    Generate a complete .docx NOD report from a JSON context dict.

    Parameters
    ----------
    context : dict
        Report context — same structure used by the HTML template.
        Must contain either:
          - context["report_meta"] + context["sections"] (generic format), OR
          - context["report_meta"] + context["charts"] / context["devices"] etc. (backend format)
    output_path : str | Path
        Destination .docx file path.

    Returns
    -------
    Path
        The output file path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    report_meta = context.get("report_meta", {})
    report_type = report_meta.get("report_type", "R-XX")

    logger.info("Generating DOCX report: type=%s → %s", report_type, output_path)

    # ── Create Document ──
    doc = Document()

    # ── Page Setup ──
    _setup_page(doc)

    # ── Default Font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_FAMILY
    font.size = FONT_BODY
    font.color.rgb = COLOR_TEXT

    # Configure heading styles
    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = FONT_FAMILY
        hs.font.color.rgb = COLOR_TEXT

    # ── Header & Footer ──
    _setup_header_footer(doc, context)

    # ── Title Page ──
    _build_title_page(doc, context)

    # ── Generic Sections (from context["sections"]) ──
    sections = context.get("sections", {})
    if sections:
        _build_generic_sections(doc, sections)

    # ── Report-Type Specific Sections ──
    report_builders = {
        "R-01": _build_r01_traffic_flow,
        "R-02": _build_r02_resource_usage,
        "R-03": _build_r03_vpn_users,
        "R-04": _build_r04_sdwan_sla,
        "R-05": _build_r05_traffic_inbound,
        "R-06": _build_r06_traffic_internal,
        "R-07": _build_r07_executive_summary,
    }

    if report_type in report_builders:
        _add_page_break(doc)
        report_builders[report_type](doc, context)

    # ── R-08: All-in-One (build all sections) ──
    if report_type == "R-08":
        for rt in ("R-01", "R-02", "R-03", "R-04", "R-05", "R-06", "R-07"):
            if rt in report_builders:
                _add_page_break(doc)
                report_builders[rt](doc, context)
        _build_classification_page(doc, context)

    # ── Save ──
    doc.save(str(output_path))
    logger.info("DOCX report saved: %s", output_path)
    return output_path


# =========================================================================
# CLI INTERFACE
# =========================================================================

def main():
    """
    CLI entry point:
      python generate_docx.py <context.json> [output.docx]

    Reads a JSON context file and produces a .docx report.
    """
    if len(sys.argv) < 2:
        print("Usage: python generate_docx.py <context.json> [output.docx]")
        print()
        print("Arguments:")
        print("  context.json   Path to the JSON context file")
        print("  output.docx    Output DOCX path (default: <context_stem>.docx)")
        sys.exit(1)

    context_path = Path(sys.argv[1])
    if not context_path.exists():
        print(f"Error: Context file not found: {context_path}")
        sys.exit(1)

    # Default output: same stem as input, .docx extension
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = context_path.with_suffix(".docx")

    # Load context
    with open(context_path, "r", encoding="utf-8") as f:
        context = json.load(f)

    # Generate
    result = generate_docx_report(context, output_path)
    print(f"✓ Report generated: {result}")


if __name__ == "__main__":
    main()
